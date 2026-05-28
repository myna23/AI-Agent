"""
Zambia GeoHub AI — Single intelligent chat interface.

The AI detects intent from the user's message:
  - Questions / exploration  → answer with map
  - "generate a report"      → build report, offer Word + PDF download
  - "summarise / summarize"  → plain-language dataset summary
  - "what data is available" → list the catalog

Run locally:
    streamlit run app.py
"""

import streamlit as st
from streamlit_folium import st_folium

from hub.client import HubClient
import hub.client as _hub_client_module
from ai.model_client import ModelClient, PROVIDERS, BEST_MODELS, DEFAULT_PROVIDER, DEFAULT_MODEL, fetch_available_models
from ai.prompts import (
    chatbot_system_prompt,
    chatbot_user_prompt,
    summarizer_system_prompt,
    summarizer_prompt,
    report_system_prompt,
    report_prompt,
    map_analysis_system_prompt,
    map_analysis_user_prompt,
)
from reports.builder import ReportBuilder
from utils.geo_utils import (
    make_folium_map, summarize_geojson, geojson_to_sample_rows,
    haversine_km, features_within_km, polygon_centroid, assign_districts,
    _point_in_polygon,
)
import json as _json_mod
import os as _os_mod
import base64 as _base64_mod
import io as _io_mod
import math

# ---------------------------------------------------------------------------
# Toolbar icon SVGs — inline, rendered directly in HTML (no JS/CSS tricks)
# ---------------------------------------------------------------------------
_S = 'width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#888" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"'
_TB_ICONS = [
    # 0 pencil
    f'<svg {_S}><path d="M11 4H4a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h14a2 2 0 0 0 2-2v-7"/><path d="M18.5 2.5a2.121 2.121 0 0 1 3 3L12 15l-4 1 1-4 9.5-9.5z"/></svg>',
    # 1 refresh
    f'<svg {_S}><polyline points="1 4 1 10 7 10"/><path d="M3.51 15a9 9 0 1 0 2.13-9.36L1 10"/></svg>',
    # 2 copy (two squares)
    f'<svg {_S}><rect x="9" y="9" width="13" height="13" rx="2"/><path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"/></svg>',
    # 3 download / save
    f'<svg {_S}><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>',
    # 4 trash
    f'<svg {_S}><polyline points="3 6 5 6 21 6"/><path d="M19 6l-1 14a2 2 0 0 1-2 2H8a2 2 0 0 1-2-2L5 6"/><path d="M10 11v6"/><path d="M14 11v6"/><path d="M9 6V4a1 1 0 0 1 1-1h4a1 1 0 0 1 1 1v2"/></svg>',
    # 5 thumbs up
    f'<svg {_S}><path d="M14 9V5a3 3 0 0 0-3-3l-4 9v11h11.28a2 2 0 0 0 2-1.7l1.38-9a2 2 0 0 0-2-2.3H14z"/><path d="M7 22H4a2 2 0 0 1-2-2v-7a2 2 0 0 1 2-2h3"/></svg>',
    # 6 thumbs down
    f'<svg {_S}><path d="M10 15v4a3 3 0 0 0 3 3l4-9V2H5.72a2 2 0 0 0-2 1.7l-1.38 9a2 2 0 0 0 2 2.3H10z"/><path d="M17 2h2.67A2.31 2.31 0 0 1 22 4v7a2.31 2.31 0 0 1-2.33 2H17"/></svg>',
]

# Module-level buffer for capturing partial streamed responses across reruns.
# Keyed by a session identifier derived from session_state id.
_STREAM_BUFFERS: dict = {}

# Pre-load context layers once (district boundaries + roads) for map background.
# These are shown underneath point datasets so users can see which district/road
# is nearest to each church, market, health facility, school, etc.
@st.cache_resource(show_spinner=False)
def _load_context_layers():
    data_dir = _os_mod.path.join(_os_mod.path.dirname(__file__), "data")
    layers = []
    _dist_path = _os_mod.path.join(data_dir, "districts.json")
    if _os_mod.path.exists(_dist_path):
        with open(_dist_path) as f:
            layers.append({"geojson": _json_mod.load(f), "name": "Districts", "type": "boundary"})
    _road_path = _os_mod.path.join(data_dir, "roads.json")
    if _os_mod.path.exists(_road_path):
        with open(_road_path) as f:
            layers.append({"geojson": _json_mod.load(f), "name": "Major Roads", "type": "road"})
    return layers

@st.cache_resource(show_spinner=False)
def _load_water_layer():
    """Load wetlands/lakes polygon layer for use as water-body context."""
    data_dir = _os_mod.path.join(_os_mod.path.dirname(__file__), "data")
    path = _os_mod.path.join(data_dir, "wetlands.json")
    if _os_mod.path.exists(path):
        with open(path) as f:
            return {"geojson": _json_mod.load(f), "name": "Lakes & Wetlands", "type": "water"}
    return None

_CONTEXT_LAYERS = _load_context_layers()
_WATER_LAYER = _load_water_layer()


# ---------------------------------------------------------------------------
# Language support — suffix appended to AI prompts
# ---------------------------------------------------------------------------
_LANG_INSTRUCTIONS = {
    "English":  "",
    "Nyanja":   " Please respond in Chichewa/Nyanja language.",
    "Bemba":    " Please respond in Bemba language.",
    "Tonga":    " Please respond in Tonga language.",
}

# ---------------------------------------------------------------------------
# Export full chat history as a Word document
# ---------------------------------------------------------------------------
def _export_chat_docx(messages: list) -> bytes:
    """Convert chat history to a formatted Word document and return bytes."""
    import io as _io
    from docx import Document as _Doc
    from docx.shared import Pt, RGBColor
    _doc = _Doc()
    _doc.add_heading("Zambia GeoHub AI — Chat Export", 0)
    import datetime as _dt2
    _doc.add_paragraph(f"Exported: {_dt2.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    _doc.add_paragraph()
    for m in messages:
        role = m.get("role", "")
        content = m.get("content", "")
        if role == "user":
            p = _doc.add_paragraph()
            r = p.add_run("You:  ")
            r.bold = True
            r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
            p.add_run(content)
        elif role == "assistant":
            p = _doc.add_paragraph()
            r = p.add_run("Assistant:  ")
            r.bold = True
            r.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
            p.add_run(content[:4000] if len(content) > 4000 else content)
            # Data source citation if present
            if m.get("data_source_url"):
                src_p = _doc.add_paragraph()
                src_r = src_p.add_run(f"  Source: {m.get('ds_name','Dataset')} — {m['data_source_url']}")
                src_r.font.size = Pt(8)
                src_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        _doc.add_paragraph()
    buf = _io.BytesIO()
    _doc.save(buf)
    return buf.getvalue()


def _map(geojson: dict, name: str, with_context: bool = False, highlight_location: str = "", draw_bbox: dict = None) -> object:
    """Wrapper: adds district + road context layers for point datasets."""
    ctx = _CONTEXT_LAYERS if with_context else None
    return make_folium_map(geojson, name, context_layers=ctx, highlight_location=highlight_location, draw_bbox=draw_bbox)


def _build_suggestions(question: str, has_location: bool, has_data: bool, ds_name: str) -> list:
    """Return up to 3 context-aware follow-up suggestion strings."""
    suggestions = []
    q = question.lower()
    location = question.split(" in ")[-1].strip().rstrip("?") if " in " in q else ""

    # Dataset-specific suggestions
    if "health" in ds_name.lower() or "hospital" in q or "clinic" in q:
        if has_location and location:
            suggestions.append(f"Which district in {location} has the most health facilities?")
        if "primary" not in q:
            suggestions.append("Compare primary vs secondary health facilities")
        if has_location and "school" not in q:
            suggestions.append(f"How many schools are in {location}?" if location else "How many schools are in Lusaka?")

    elif "school" in ds_name.lower() or "school" in q or "education" in q:
        if has_location and location:
            suggestions.append(f"Which district in {location} has the most schools?")
        if "primary" not in q:
            suggestions.append("Compare primary vs secondary schools")
        if has_location:
            suggestions.append(f"How many health facilities are in {location}?" if location else "What are the health facilities near schools?")

    elif "settlement" in ds_name.lower() or "settlement" in q or "village" in q:
        if has_location and location:
            suggestions.append(f"What is the flood risk for settlements in {location}?")
            suggestions.append(f"How many schools are in {location}?")
        suggestions.append("Which province has the most settlements?")

    elif "road" in ds_name.lower() or "road" in q:
        if has_location and location:
            suggestions.append(f"What is the road surface breakdown in {location}?")
        suggestions.append("Which roads are unpaved or in poor condition?")
        if has_location and location:
            suggestions.append(f"How many settlements are accessible in {location}?")

    elif "flood" in ds_name.lower() or "flood" in q:
        if has_location and location:
            suggestions.append(f"How many settlements are in flood-prone areas in {location}?")
        suggestions.append("Which province has the most flood-prone districts?")
        if has_location and location:
            suggestions.append(f"What are the roads like in {location}?")

    # Generic fallbacks if no specific suggestions added yet
    if len(suggestions) < 2:
        if has_location and location and "compare" not in q:
            suggestions.append(f"Compare districts in {location}")
        if has_data and "table" not in q:
            suggestions.append(f"Show me a breakdown of {ds_name} by district")
        if has_data and "report" not in q:
            suggestions.append(f"Generate a report on {ds_name}")
        if has_location and "flood" not in q and "risk" not in q:
            suggestions.append(f"What is the flood risk in {location or 'this area'}?")
        if has_location and "settlement" not in q:
            suggestions.append(f"How many settlements are in {location or 'this area'}?")

    return suggestions[:3]


def _suggestion_chips(question: str, has_location: bool, has_data: bool, ds_name: str, key_prefix: str = ""):
    """Follow-up suggestions — right-aligned, stacked vertically below the answer."""
    suggestions = _build_suggestions(question, has_location, has_data, ds_name)
    if not suggestions:
        return
    st.markdown("<div style='margin-top:6px'></div>", unsafe_allow_html=True)
    _spacer, _sug_col = st.columns([3, 7])
    with _sug_col:
        for j, sug in enumerate(suggestions):
            if st.button(sug, key=f"{key_prefix}_sug_{j}", use_container_width=True):
                st.session_state.messages.append({"role": "user", "content": sug})
                st.session_state._pending_question = sug
                st.rerun()


def _render_ondemand_panel(msg_idx: int, msg: dict, ctx_layers: list = None):
    """
    Show Map / Table / Chart buttons after an answer.
    Only renders the component when the user explicitly clicks the button.
    State is stored in the message dict so it survives reruns.
    """
    has_geojson = bool(msg.get("geojson") and msg["geojson"].get("features"))
    has_data = bool(msg.get("sample_features"))
    if not has_geojson and not has_data:
        return

    # Toggle buttons — label changes to hide when already shown
    _ba, _bb, _bc, _bz = st.columns([1.4, 1.4, 1.4, 8])
    with _ba:
        if has_geojson:
            if st.button("🗺️ Map", key=f"mapbtn_{msg_idx}", use_container_width=True):
                msg["map_shown"] = not msg.get("map_shown", False)
                st.rerun()
    with _bb:
        if has_data:
            if st.button("📊 Table", key=f"tblbtn_{msg_idx}", use_container_width=True):
                msg["table_shown"] = not msg.get("table_shown", False)
                st.rerun()
    with _bc:
        if has_data:
            if st.button("📈 Chart", key=f"chtbtn_{msg_idx}", use_container_width=True):
                msg["chart_shown"] = not msg.get("chart_shown", False)
                st.rerun()

    # Render requested components
    if msg.get("map_shown") and has_geojson:
        gjson = msg["geojson"]
        ds_name_lower = (msg.get("ds_name", "") or "").lower()
        layers = ctx_layers if ctx_layers else _CONTEXT_LAYERS
        # For non-point datasets (polygons/lines), use water layer as context for dam datasets
        if _is_point_geojson(gjson):
            map_layers = layers
        elif ("dam" in ds_name_lower or "reservoir" in ds_name_lower) and _WATER_LAYER:
            map_layers = [_WATER_LAYER]
        else:
            map_layers = None
        st_folium(
            make_folium_map(
                gjson,
                msg.get("ds_name", ""),
                context_layers=map_layers,
                highlight_location=msg.get("location", ""),
                buffer_center=msg.get("buffer_center"),
                buffer_radius_km=msg.get("buffer_radius_km"),
                buffer_label=(
                    (f"{msg.get('buffer_radius_km')} km radius — {msg.get('location')}"
                     if msg.get("location") else
                     f"{msg.get('buffer_radius_km')} km radius")
                    if msg.get("buffer_radius_km") else ""
                ),
                draw_bbox=msg.get("draw_bbox"),
            ),
            width=720, height=340, returned_objects=[], key=f"map_{msg_idx}"
        )

    if msg.get("table_shown") and has_data:
        _render_data_tables(msg["sample_features"], msg.get("ds_name", "Data"), key_prefix=f"tbl_{msg_idx}")

    if msg.get("chart_shown") and has_data:
        _render_charts_only(msg["sample_features"], msg.get("ds_name", "Data"), key_prefix=f"cht_{msg_idx}")


def _render_location_overview(location: str, loc_type: str, hub, key_prefix: str = ""):
    """
    Render a compact 'Area Overview' card for a district or province.

    Fetches live counts from the main thematic datasets and displays them
    as metric tiles so users instantly understand what's in an area they
    may not be familiar with.

    Datasets queried:
      - Health facilities (GRID3)
      - Schools (GRID3)
      - Points of Interest — Commercial, Religion, Farm subcategories
      - Settlements (GRID3)
    """
    if not location:
        return

    # Dataset URLs from catalog — picked from seed catalog for reliability
    _OV_DATASETS = [
        {
            "label": "Health Facilities",
            "emoji": "🏥",
            "url_key": "GRID3_ZMB_HealthFac",
            "color": "#c1121f",
        },
        {
            "label": "Schools",
            "emoji": "🏫",
            "url_key": "GRID3_ZMB_School",
            "color": "#1d3557",
        },
        {
            "label": "Settlements",
            "emoji": "🏘️",
            "url_key": "GRID3_Zambia_Operational_Settlement_Points_and_Names",
            "color": "#7b4f12",
        },
        {
            "label": "Points of Interest",
            "emoji": "📍",
            "url_key": "GRID3_Zambia_Operational_Points_of_Interest",
            "color": "#2d6a4f",
        },
    ]

    # Find URLs from catalog
    catalog = hub.get_catalog()
    _url_map = {}
    for entry in catalog:
        for ds in _OV_DATASETS:
            if ds["url_key"] in entry.get("url", ""):
                if ds["label"] not in _url_map:
                    _url_map[ds["label"]] = entry["url"]

    # Use static file fallback URLs for robustness
    from hub.client import _STATIC_MAP as _SM, _DATA_DIR as _DD
    import os as _os, json as _json
    _static_fallbacks = {}
    for ds in _OV_DATASETS:
        for key, fname in _SM.items():
            if ds["url_key"] in key or key in ds["url_key"]:
                path = _os.path.join(_DD, fname)
                if _os.path.exists(path):
                    _static_fallbacks[ds["label"]] = path

    is_province = (loc_type == "province")
    filter_field = "Province" if is_province else "District"

    counts = {}
    poi_breakdown = {}

    for ds in _OV_DATASETS:
        label = ds["label"]
        count = None

        # Try live count first
        url = _url_map.get(label)
        if url:
            try:
                if is_province:
                    count = hub.count_features(url, province_filter=location)
                else:
                    count = hub.count_features(url, district_filter=location)
            except Exception:
                count = None

        # Fallback: count from static file
        if count is None and label in _static_fallbacks:
            try:
                with open(_static_fallbacks[label]) as _f:
                    _gj = _json.load(_f)
                _feats = _gj.get("features", [])
                _filtered = [
                    f for f in _feats
                    if location.lower() in (
                        (f.get("properties") or {}).get(filter_field, "") or
                        (f.get("properties") or {}).get(filter_field.upper(), "") or ""
                    ).lower()
                ]
                count = len(_filtered)

                # For POI, also break down by Type
                if label == "Points of Interest" and _filtered:
                    for f in _filtered:
                        t = (f.get("properties") or {}).get("Type", "Other") or "Other"
                        poi_breakdown[t] = poi_breakdown.get(t, 0) + 1
            except Exception:
                count = None

        if count is not None:
            counts[label] = count

    if not counts:
        return

    st.markdown(f"#### 📊 {location} — Area Overview")
    st.caption(
        f"Key infrastructure counts for **{location}** {'Province' if is_province else 'District'} "
        f"(from live GeoHub data or pre-loaded sample):"
    )

    _cols = st.columns(len(_OV_DATASETS))
    for i, ds in enumerate(_OV_DATASETS):
        label = ds["label"]
        if label in counts:
            _cols[i].metric(
                label=f"{ds['emoji']} {label}",
                value=f"{counts[label]:,}",
            )

    # POI breakdown by type if available
    if poi_breakdown:
        top_types = sorted(poi_breakdown.items(), key=lambda x: x[1], reverse=True)[:6]
        _type_md = "  ".join(f"**{t}**: {c}" for t, c in top_types)
        st.caption(f"POI breakdown — {_type_md}")

    st.markdown("---")


def _render_charts_only(sample_features: list, ds_name: str, key_prefix: str = ""):
    """Bar charts for categorical fields; line/bar chart for numeric fields as fallback."""
    import pandas as _pd
    rows = [r for r in sample_features if "_note" not in r]
    if not rows:
        st.info("No data available for chart.")
        return
    df = _pd.DataFrame(rows)
    # Drop non-useful columns
    drop_cols = [c for c in df.columns if c.lower() in ("objectid", "fid", "globalid", "geometry", "shape_area", "shape_length")]
    df = df.drop(columns=[c for c in drop_cols if c in df.columns])

    shown = 0
    # 1. Categorical bar charts
    cat_fields = ["District", "DISTRICT", "Province", "PROVINCE", "Type", "TYPE",
                  "SubType", "Facility_T", "fclass", "surface", "Status", "STATUS",
                  "Classifica", "S_CLASS", "DOMINANT"]
    for field in cat_fields:
        if field not in df.columns or shown >= 2:
            continue
        counts = df[field].dropna().astype(str)
        counts = counts[counts != "None"].value_counts().head(15)
        if len(counts) < 2:
            continue
        chart_df = counts.reset_index()
        chart_df.columns = [field, "Count"]
        st.markdown(f"**{ds_name} — by {field}**")
        st.bar_chart(chart_df.set_index(field)["Count"])
        shown += 1

    # 2. Numeric bar charts as fallback if no categoricals found
    if shown == 0:
        skip = {"lat", "long", "lat_", "long_", "x", "y", "objectid", "fid"}
        num_cols = [c for c in df.select_dtypes(include="number").columns if c.lower() not in skip][:3]
        for col in num_cols:
            series = df[col].dropna()
            if len(series) < 2:
                continue
            # Use index label if available, else row number
            label_col = next((c for c in ["Name", "NAME", "District", "DISTRICT", "Province"] if c in df.columns), None)
            chart_df = df[[label_col, col]].dropna() if label_col else series.reset_index()
            if label_col:
                chart_df = chart_df.set_index(label_col)
            st.markdown(f"**{ds_name} — {col}**")
            st.bar_chart(chart_df[col] if label_col else chart_df)
            shown += 1

    if shown == 0:
        st.info("No suitable fields found for a chart with this dataset.")


def _render_data_tables(sample_features: list, ds_name: str, key_prefix: str = ""):
    """
    Render an expandable data table + bar charts for sample feature records.
    - Shows all sample rows as a scrollable dataframe
    - For categorical fields (District, Province, Type, etc.) shows a bar chart
    """
    import pandas as _pd

    if not sample_features:
        return
    # Filter out internal placeholder notes
    rows = [r for r in sample_features if "_note" not in r]
    if not rows:
        return

    df = _pd.DataFrame(rows)

    # Drop geometry columns if any snuck in
    geo_cols = [c for c in df.columns if c.lower() in ("geometry", "shape", "shape_area", "shape_length", "objectid", "fid", "globalid")]
    df = df.drop(columns=[c for c in geo_cols if c in df.columns])

    with st.expander(f"Data table — {ds_name} ({len(rows)} records)", expanded=True):
        st.dataframe(df, use_container_width=True, height=260)

        # CSV download
        st.download_button(
            "⬇️ Download CSV",
            df.to_csv(index=False).encode("utf-8"),
            file_name=f"{ds_name.replace(' ', '_')}_data.csv",
            mime="text/csv",
            key=f"{key_prefix}_csv",
            use_container_width=False,
        )

        # Auto summary stats for numeric columns
        num_cols = df.select_dtypes(include="number").columns.tolist()
        num_cols = [c for c in num_cols if c.upper() not in ("OBJECTID", "FID", "LAT", "LONG", "LAT_", "LONG_", "X", "Y")]
        if num_cols:
            st.markdown("**Summary statistics**")
            st.dataframe(df[num_cols].describe().round(2), use_container_width=True)

        # Value counts for key categorical columns
        cat_priority = ["District", "DISTRICT", "Province", "PROVINCE", "Type", "TYPE",
                        "SubType", "Facility_T", "fclass", "surface", "Status", "STATUS"]
        shown_cat = 0
        for field in cat_priority:
            if field not in df.columns or shown_cat >= 2:
                continue
            counts = df[field].dropna().astype(str)
            counts = counts[counts != "None"].value_counts().head(10)
            if len(counts) < 2:
                continue
            st.markdown(f"**Breakdown by {field}**")
            count_df = counts.reset_index()
            count_df.columns = [field, "Count"]
            count_df["% of total"] = (count_df["Count"] / count_df["Count"].sum() * 100).round(1).astype(str) + "%"
            st.dataframe(count_df, use_container_width=True, hide_index=True)
            shown_cat += 1


def _is_point_geojson(geojson: dict) -> bool:
    feats = geojson.get("features", [])
    if not feats:
        return False
    g = feats[0].get("geometry")
    return bool(g and g.get("type") == "Point")

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="Zambia GeoHub AI",
    page_icon="🗺️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ---------------------------------------------------------------------------
# Styling — clean, Hub-like look
# ---------------------------------------------------------------------------
st.markdown("""
<style>
/* ── Global font & base ─────────────────────────────────────────────── */
html, body, [class*="css"] { font-family: 'Inter', 'Segoe UI', sans-serif; }

/* ── Hero header banner ─────────────────────────────────────────────── */
.zmb-hero {
    background: linear-gradient(135deg, #1d3557 0%, #2a6496 60%, #1a6b3c 100%);
    border-radius: 16px;
    padding: 28px 36px 22px 36px;
    margin-bottom: 4px;
    color: white;
}
.zmb-hero h1 {
    margin: 0 0 6px 0;
    font-size: 2rem;
    font-weight: 700;
    letter-spacing: -0.5px;
    color: white !important;
}
.zmb-hero p {
    margin: 0;
    font-size: 0.95rem;
    opacity: 0.85;
    line-height: 1.5;
}
.zmb-hero-badge {
    display: inline-block;
    background: rgba(255,255,255,0.18);
    border: 1px solid rgba(255,255,255,0.3);
    border-radius: 20px;
    font-size: 0.75rem;
    font-weight: 600;
    padding: 3px 12px;
    margin-right: 6px;
    margin-top: 10px;
}

/* ── Suggestion chips ───────────────────────────────────────────────── */
.zmb-chips { display: flex; flex-wrap: wrap; gap: 8px; margin: 16px 0 4px 0; }
.zmb-chip {
    background: #f0f4f9;
    border: 1px solid #d0dbe8;
    border-radius: 20px;
    padding: 6px 16px;
    font-size: 0.82rem;
    color: #1d3557;
    cursor: pointer;
    transition: all 0.15s;
    font-weight: 500;
}
.zmb-chip:hover { background: #1d3557; color: white; border-color: #1d3557; }

/* ── Feature cards (welcome grid) ───────────────────────────────────── */
.zmb-card-grid { display: flex; gap: 12px; flex-wrap: wrap; margin: 12px 0; }
.zmb-card {
    flex: 1; min-width: 180px;
    background: white;
    border: 1px solid #e4eaf2;
    border-radius: 12px;
    padding: 16px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.05);
}
.zmb-card-icon { font-size: 1.6rem; margin-bottom: 8px; }
.zmb-card-title { font-weight: 600; font-size: 0.88rem; color: #1d3557; margin-bottom: 4px; }
.zmb-card-desc { font-size: 0.78rem; color: #5a6a7a; line-height: 1.4; }

/* ── Chat messages ──────────────────────────────────────────────────── */
[data-testid="stChatMessage"] {
    border-radius: 12px !important;
    padding: 4px 0 !important;
}

/* ── Floating chat bubble (Hub iframe embed) ────────────────────────── */
#zmb-chat-btn {
    position: fixed; bottom: 28px; right: 28px;
    width: 58px; height: 58px; border-radius: 50%;
    background: #1d3557; color: white; font-size: 26px;
    border: none; cursor: pointer;
    box-shadow: 0 4px 16px rgba(0,0,0,0.25);
    z-index: 9999; display: flex;
    align-items: center; justify-content: center;
}
#zmb-chat-btn:hover { background: #457b9d; }
#zmb-chat-panel {
    position: fixed; bottom: 100px; right: 28px;
    width: 380px; height: 520px; background: white;
    border-radius: 16px; box-shadow: 0 8px 32px rgba(0,0,0,0.2);
    z-index: 9998; display: none; flex-direction: column;
    overflow: hidden; border: 1px solid #e0e0e0;
}
#zmb-chat-panel.open { display: flex; }
#zmb-chat-header {
    background: #1d3557; color: white;
    padding: 14px 18px; font-weight: 600; font-size: 15px;
    display: flex; justify-content: space-between; align-items: center;
}
#zmb-chat-close { cursor: pointer; font-size: 20px; background: none; border: none; color: white; }
#zmb-chat-body { flex: 1; overflow-y: auto; padding: 14px; font-size: 13px; background: #f8fbfd; }
#zmb-chat-footer { padding: 10px; border-top: 1px solid #eee; display: flex; gap: 8px; background: white; }
#zmb-chat-input { flex: 1; border: 1px solid #ccc; border-radius: 8px; padding: 8px 12px; font-size: 13px; outline: none; }
#zmb-chat-send { background: #1d3557; color: white; border: none; border-radius: 8px; padding: 8px 14px; cursor: pointer; font-size: 13px; }
#zmb-chat-send:hover { background: #457b9d; }
.zmb-msg-user { background: #1d3557; color: white; border-radius: 12px 12px 2px 12px; padding: 8px 12px; margin: 6px 0 6px 30px; font-size: 13px; }
.zmb-msg-ai { background: white; border: 1px solid #dde; border-radius: 12px 12px 12px 2px; padding: 8px 12px; margin: 6px 30px 6px 0; font-size: 13px; }

/* ── Action toolbar ─────────────────────────────────────────────────── */
.zmb-tb { display: none; }

/* Each toolbar column: relative so the SVG overlay can be absolute inside it */
[data-testid="stColumn"]:has(.zmb-ic) {
    position: relative !important;
    overflow: visible !important;
}
/* Zero-height wrapper — the SVG markdown block takes no vertical space */
.element-container:has(.zmb-ic) {
    height: 0 !important;
    overflow: visible !important;
    padding: 0 !important;
    margin: 0 !important;
    line-height: 0 !important;
    position: static !important;
}
/* SVG overlay: sits behind the button, centered in the column */
.zmb-ic {
    position: absolute !important;
    top: 0 !important;
    left: 0 !important;
    right: 0 !important;
    bottom: 0 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    pointer-events: none !important;
    z-index: 1 !important;
}
.zmb-ic svg { display: block !important; }

/* Toolbar buttons — fully invisible, just the click zone */
[data-testid="stColumn"]:has(.zmb-ic) button {
    background: transparent !important;
    background-color: transparent !important;
    color: transparent !important;
    font-size: 0 !important;
    border: none !important;
    border-color: transparent !important;
    box-shadow: none !important;
    outline: none !important;
    position: relative !important;
    z-index: 5 !important;
}
[data-testid="stColumn"]:has(.zmb-ic) button:hover {
    background-color: rgba(0,0,0,0.07) !important;
    border-radius: 6px !important;
}
[data-testid="stColumn"]:has(.zmb-ic) button:focus,
[data-testid="stColumn"]:has(.zmb-ic) button:active {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
}
/* Liked/disliked confirmation tick */
.zmb-tb-liked {
    color: #4caf50; font-size: 14px; text-align: center; line-height: 32px;
}

/* ── Intent badge ───────────────────────────────────────────────────── */
.intent-badge {
    display: inline-block; font-size: 11px; font-weight: 600;
    padding: 2px 10px; border-radius: 20px; margin-bottom: 6px;
}
.intent-chat    { background: #e8f4fd; color: #1d3557; }
.intent-report  { background: #e8f8f0; color: #1a6b3c; }
.intent-summary { background: #fff4e6; color: #7a4800; }

/* ── Sidebar polish ─────────────────────────────────────────────────── */
[data-testid="stSidebar"] {
    background: #0e1a2b !important;
    border-right: 1px solid #1d3557;
}
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] div.stMarkdown,
[data-testid="stSidebar"] .stCaption {
    color: #e8edf3 !important;
}
[data-testid="stSidebar"] h3 {
    color: #ffffff !important;
    font-size: 1.1rem;
    font-weight: 700;
    letter-spacing: 0.3px;
}
[data-testid="stSidebar"] h4 {
    color: #a8c4e0 !important;
    font-size: 0.82rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    margin-top: 4px;
}
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stTextInput label {
    color: #a8c4e0 !important;
}
[data-testid="stSidebar"] [data-baseweb="select"] > div {
    background: #1a2d45 !important;
    border-color: #2d4a6a !important;
    color: #e8edf3 !important;
}
[data-testid="stSidebar"] input {
    background: #1a2d45 !important;
    border-color: #2d4a6a !important;
    color: #e8edf3 !important;
}
[data-testid="stSidebar"] input::placeholder {
    color: #7a9bbf !important;
    opacity: 1 !important;
}
[data-testid="stSidebar"] textarea::placeholder {
    color: #7a9bbf !important;
    opacity: 1 !important;
}
/* File uploader drop zone — dark background, clear text */
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
    background-color: #1a2d45 !important;
    border: 1.5px dashed #4a7fa8 !important;
    border-radius: 6px !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] span,
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] p {
    color: #d0e4f5 !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] small {
    color: #8ab4d4 !important;
}
[data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button {
    background-color: #2d5a87 !important;
    color: #ffffff !important;
    border: 1px solid #4a7fa8 !important;
}
/* Ensure folium map iframe is visible */
[data-testid="stSidebar"] iframe {
    display: block !important;
    visibility: visible !important;
    opacity: 1 !important;
    min-height: 320px !important;
    height: 320px !important;
}
[data-testid="stSidebar"] [data-testid="stCustomComponentV1"] {
    min-height: 320px !important;
    height: 320px !important;
    display: block !important;
}
/* folium / st_folium in main content area */
[data-testid="stMain"] [data-testid="stCustomComponentV1"],
[data-testid="stMain"] iframe {
    display: block !important;
    min-height: 480px !important;
}
[data-testid="stSidebar"] .element-container:has(iframe) {
    min-height: 320px !important;
    height: 320px !important;
}
[data-testid="stSidebar"] hr {
    border-color: #1d3557 !important;
}
/* Expander header — dark background, light text */
[data-testid="stSidebar"] [data-testid="stExpander"] {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
}
[data-testid="stSidebar"] [data-testid="stExpander"] summary,
[data-testid="stSidebar"] [data-testid="stExpander"] summary:hover {
    background: transparent !important;
    color: #e8edf3 !important;
}
/* Expander content area — remove white box */
[data-testid="stSidebar"] [data-testid="stExpander"] > div[data-testid="stExpanderDetails"],
[data-testid="stSidebar"] details > div,
[data-testid="stSidebar"] details[open] > div {
    background: transparent !important;
    background-color: transparent !important;
    border: none !important;
    box-shadow: none !important;
}
[data-testid="stSidebar"] details {
    background: transparent !important;
    border: none !important;
}
/* Primary button (New Chat) — keep as a real button */
[data-testid="stSidebar"] [data-testid="baseButton-primary"] {
    background: #1d3557 !important;
    color: white !important;
    border: 1px solid #2d4a6a !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
}
[data-testid="stSidebar"] [data-testid="baseButton-primary"]:hover {
    background: #2a6496 !important;
}
/* Secondary buttons (chat history items) — transparent, plain text */
[data-testid="stSidebar"] button[kind="secondary"],
[data-testid="stSidebar"] button[data-testid="baseButton-secondary"],
[data-testid="stSidebar"] [data-testid="baseButton-secondary"],
[data-testid="stSidebar"] [data-testid="stButton"] button {
    background: transparent !important;
    background-color: transparent !important;
    background-image: none !important;
    border: none !important;
    box-shadow: none !important;
    outline: none !important;
    color: #c8dff0 !important;
    font-weight: 400 !important;
    font-size: 0.83rem !important;
    text-align: left !important;
    justify-content: flex-start !important;
    padding: 0.2rem 0.5rem !important;
    border-radius: 4px !important;
}
[data-testid="stSidebar"] button[kind="secondary"]:hover,
[data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:hover,
[data-testid="stSidebar"] [data-testid="stButton"] button:hover {
    background: rgba(255,255,255,0.07) !important;
    background-color: rgba(255,255,255,0.07) !important;
    color: #e8f4ff !important;
}
[data-testid="stSidebar"] button[kind="secondary"]:focus,
[data-testid="stSidebar"] button[kind="secondary"]:active,
[data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:focus,
[data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:active {
    background: transparent !important;
    background-color: transparent !important;
    box-shadow: none !important;
}
[data-testid="stSidebar"] [data-testid="stButton"] {
    background: transparent !important;
    background-color: transparent !important;
}
[data-testid="stSidebar"] [data-testid="stButton"] button p,
[data-testid="stSidebar"] button[kind="secondary"] p {
    white-space: nowrap !important;
    overflow: hidden !important;
    text-overflow: ellipsis !important;
    max-width: 100% !important;
    color: #c8dff0 !important;
}
/* Recents label — plain text, no background */
.zmb-recents-label {
    color: #7a9bbf !important;
    font-size: 0.72rem !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.8px !important;
    margin: 8px 0 4px 8px !important;
    background: transparent !important;
    padding: 0 !important;
}
/* Ensure no white box appears on caption-like elements in sidebar */
[data-testid="stSidebar"] .stMarkdown p {
    background: transparent !important;
}
/* Topic category buttons on welcome screen */
[data-testid="stHorizontalBlock"] [data-testid="baseButton-secondary"]:has-text {
    /* fallback — handled below */
}
/* Style the topic buttons to look like pill buttons */
div[data-testid="column"] [data-testid="baseButton-secondary"] {
    background: #1a3a5c !important;
    border: 1px solid #2d5a87 !important;
    border-radius: 24px !important;
    color: #d0e8ff !important;
    font-weight: 600 !important;
    font-size: 0.83rem !important;
    padding: 0.45rem 0.9rem !important;
    transition: background 0.15s !important;
    white-space: nowrap !important;
}
div[data-testid="column"] [data-testid="baseButton-secondary"]:hover {
    background: #2a5a8c !important;
    border-color: #4a8abf !important;
    color: #ffffff !important;
}
</style>

<button id="zmb-chat-btn" title="Ask the Zambia GeoHub AI">🗺️</button>
<div id="zmb-chat-panel">
  <div id="zmb-chat-header">
    <span>Zambia GeoHub AI</span>
    <button id="zmb-chat-close">✕</button>
  </div>
  <div id="zmb-chat-body">
    <div class="zmb-msg-ai" id="zmb-welcome">Hi! Ask me anything about Zambia's geospatial data, or say "generate a report on health facilities" or "summarise the schools dataset".</div>
  </div>
  <div id="zmb-chat-footer">
    <input id="zmb-chat-input" type="text" placeholder="Ask about Zambia data..." />
    <button id="zmb-chat-send">Send</button>
  </div>
</div>
<script>
const btn = document.getElementById('zmb-chat-btn');
const panel = document.getElementById('zmb-chat-panel');
const closeBtn = document.getElementById('zmb-chat-close');
const input = document.getElementById('zmb-chat-input');
const sendBtn = document.getElementById('zmb-chat-send');
const body = document.getElementById('zmb-chat-body');
btn.onclick = () => panel.classList.toggle('open');
closeBtn.onclick = () => panel.classList.remove('open');
function addMsg(text, role) {
    const div = document.createElement('div');
    div.className = role === 'user' ? 'zmb-msg-user' : 'zmb-msg-ai';
    div.innerText = text;
    body.appendChild(div);
    body.scrollTop = body.scrollHeight;
}
// Detect dataset context from URL params and customise welcome message
(function() {
    const params = new URLSearchParams(window.location.search);
    const dsName = params.get('dataset_name');
    if (dsName) {
        const welcome = document.getElementById('zmb-welcome');
        if (welcome) {
            welcome.innerText = 'Hi! I can see you\'re viewing the ' + dsName + ' dataset. Ask me any question about it — for example "What districts are covered?" or "Generate a report on this dataset".';
        }
    }
})();
sendBtn.onclick = () => {
    const q = input.value.trim();
    if (!q) return;
    addMsg(q, 'user');
    input.value = '';
    setTimeout(() => addMsg('Use the main chat above for full AI responses with maps and downloads.', 'ai'), 800);
};
input.addEventListener('keydown', e => { if (e.key === 'Enter') sendBtn.click(); });
</script>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Toolbar icon injector — JS MutationObserver adds .zmb-icon-N to each button
# CSS class selectors then apply the SVG background-image reliably
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Open data-source URL in new tab if a topic button was just clicked
# ---------------------------------------------------------------------------
if st.session_state.get("_open_url"):
    _url_to_open = st.session_state.pop("_open_url")
    st.markdown(
        f'<script>window.open("{_url_to_open}", "_blank");</script>',
        unsafe_allow_html=True,
    )

# ---------------------------------------------------------------------------
# Clients
# ---------------------------------------------------------------------------
@st.cache_resource(show_spinner=False)
def get_hub(_v=6): return HubClient()

@st.cache_resource(show_spinner=False)
def get_ai_client(provider: str, model: str, api_key_hash: str):
    """Cached by (provider, model, hash of key) so a new key invalidates the cache."""
    return ModelClient(provider, model, api_key_hash)

@st.cache_resource(show_spinner=False)
def get_builder(_v=3): return ReportBuilder()

hub = get_hub()
builder = get_builder()

def _resolve_ai_key(provider: str) -> str:
    """Read API key: session state first, then .env / Streamlit secrets."""
    from ai.model_client import PROVIDERS
    env_var = PROVIDERS[provider]["env_key"]
    # Session state key set by the settings UI
    ss_key = f"ai_key_{provider}"
    if st.session_state.get(ss_key, "").strip():
        return st.session_state[ss_key].strip()
    # Environment / Streamlit secrets
    import os
    key = os.getenv(env_var, "")
    if not key:
        try:
            key = st.secrets.get(env_var, "")
        except Exception:
            pass
    return key

_ai_provider = st.session_state.get("ai_provider", DEFAULT_PROVIDER)
_ai_model    = st.session_state.get("ai_model",    DEFAULT_MODEL)
_ai_key      = _resolve_ai_key(_ai_provider)
# Cache key includes the actual key value so changing it invalidates the cache
claude = get_ai_client(_ai_provider, _ai_model, _ai_key)

# ---------------------------------------------------------------------------
# Session state — must be initialised before the sidebar reads them
# ---------------------------------------------------------------------------
if "chat_sessions" not in st.session_state:
    st.session_state.chat_sessions = []

# ---------------------------------------------------------------------------
# Sidebar
# ---------------------------------------------------------------------------
import datetime as _dt

with st.sidebar:
    st.caption("Zambia GeoHub AI")

    import os as _os
    import datetime as _dt_chat
    import uuid as _uuid
    _mai_configured = bool(_os.getenv("MAI_FACTORY_TOKEN", ""))

    def _save_current_chat():
        _msgs = st.session_state.get("messages", [])
        if not _msgs:
            return
        _title = next((m["content"][:50] for m in _msgs if m["role"] == "user"), "Chat")
        _title = (_title + "…") if len(_title) == 50 else _title
        _existing_ids = [s["id"] for s in st.session_state.chat_sessions]
        _cid = st.session_state.get("_current_chat_id", str(_uuid.uuid4()))
        if _cid in _existing_ids:
            for _s in st.session_state.chat_sessions:
                if _s["id"] == _cid:
                    _s["messages"] = list(_msgs); _s["title"] = _title
        else:
            st.session_state.chat_sessions.insert(0, {
                "id": _cid, "title": _title, "messages": list(_msgs),
                "time": _dt_chat.datetime.now().strftime("%b %d, %H:%M"),
            })
        st.session_state.chat_sessions = st.session_state.chat_sessions[:20]

    # Model — at top so it's always accessible without scrolling
    if _mai_configured:
        _sb_model_opts = [
            ("Claude Sonnet", "WB mAI Factory (Claude)", "claude-sonnet-4-5"),
            ("Claude Haiku",  "WB mAI Factory (Claude)", "claude-haiku-4-5"),
            ("GPT-4o",        "WB mAI Factory (GPT)",    "gpt-4o"),
            ("GPT-4o mini",   "WB mAI Factory (GPT)",    "gpt-4o-mini"),
        ]
    else:
        _sb_model_opts = [
            ("Claude Sonnet", "Anthropic (Claude)", "claude-sonnet-4-6"),
            ("Claude Opus",   "Anthropic (Claude)", "claude-opus-4-6"),
            ("GPT-4o",        "OpenAI (GPT)",       "gpt-4o"),
            ("Gemini Flash",  "Google (Gemini)",    "gemini-2.0-flash"),
        ]
    _sb_opt_labels = [o[0] for o in _sb_model_opts]
    _sb_opt_models = [o[2] for o in _sb_model_opts]
    _sb_cur_m = st.session_state.get("ai_model", DEFAULT_MODEL)
    _sb_cur_idx = _sb_opt_models.index(_sb_cur_m) if _sb_cur_m in _sb_opt_models else 0
    _sb_sel = st.selectbox("Model", _sb_opt_labels, index=_sb_cur_idx, key="sb_model_select")
    _sb_sel_opt = _sb_model_opts[_sb_opt_labels.index(_sb_sel)]
    if _sb_sel_opt[2] != _sb_cur_m:
        st.session_state["ai_provider"] = _sb_sel_opt[1]
        st.session_state["ai_model"]    = _sb_sel_opt[2]
        st.rerun()

    # Language
    _lang = st.selectbox(
        "Language", options=list(_LANG_INSTRUCTIONS.keys()),
        index=list(_LANG_INSTRUCTIONS.keys()).index(st.session_state.get("_lang", "English")),
        key="lang_select",
    )
    st.session_state["_lang"] = _lang

    with st.expander("Compare Two Areas"):
        st.caption("Enter two districts/provinces and a topic.")
        _cmp_col1, _cmp_col2 = st.columns(2)
        with _cmp_col1:
            _cmp_a = st.text_input("Area A", placeholder="e.g. Lusaka", key="cmp_area_a")
        with _cmp_col2:
            _cmp_b = st.text_input("Area B", placeholder="e.g. Kitwe", key="cmp_area_b")
        _cmp_topic = st.text_input("Topic", placeholder="e.g. health facilities", key="cmp_topic")
        if st.button("Compare", key="cmp_btn", use_container_width=True):
            if _cmp_a.strip() and _cmp_b.strip() and _cmp_topic.strip():
                _cmp_q = f"Compare {_cmp_a.strip()} and {_cmp_b.strip()} in terms of {_cmp_topic.strip()}"
                st.session_state.messages.append({"role": "user", "content": _cmp_q})
                st.session_state._pending_question = _cmp_q
                st.rerun()
            else:
                st.warning("Fill in both areas and a topic.")

    # ------------------------------------------------------------------
    # Select an Area — collapsed expander
    # ------------------------------------------------------------------
    with st.expander("Select an Area"):

        _AREA_BBOXES = {
            # Provinces
            "Central Province":       {"min_lat":-15.8,"max_lat":-12.3,"min_lon":25.8,"max_lon":30.8},
            "Copperbelt Province":    {"min_lat":-13.8,"max_lat":-11.8,"min_lon":26.8,"max_lon":29.2},
            "Eastern Province":       {"min_lat":-15.5,"max_lat":-11.2,"min_lon":30.8,"max_lon":33.8},
            "Luapula Province":       {"min_lat":-12.8,"max_lat":-8.5, "min_lon":28.0,"max_lon":30.8},
            "Lusaka Province":        {"min_lat":-16.0,"max_lat":-14.3,"min_lon":27.5,"max_lon":29.5},
            "Muchinga Province":      {"min_lat":-13.2,"max_lat":-9.0, "min_lon":30.5,"max_lon":33.5},
            "Northern Province":      {"min_lat":-12.0,"max_lat":-7.8, "min_lon":28.5,"max_lon":32.8},
            "North-Western Province": {"min_lat":-14.2,"max_lat":-9.2, "min_lon":21.8,"max_lon":26.5},
            "Southern Province":      {"min_lat":-18.5,"max_lat":-15.2,"min_lon":25.2,"max_lon":29.8},
            "Western Province":       {"min_lat":-18.2,"max_lat":-13.0,"min_lon":21.3,"max_lon":25.8},
            # Major towns & districts
            "Chililabombwe":          {"min_lat":-12.42,"max_lat":-12.28,"min_lon":27.78,"max_lon":27.98},
            "Chingola":               {"min_lat":-12.58,"max_lat":-12.42,"min_lon":27.78,"max_lon":28.02},
            "Chipata":                {"min_lat":-13.70,"max_lat":-13.48,"min_lon":32.50,"max_lon":32.80},
            "Choma":                  {"min_lat":-16.92,"max_lat":-16.72,"min_lon":26.88,"max_lon":27.12},
            "Kabwe":                  {"min_lat":-14.50,"max_lat":-14.30,"min_lon":28.28,"max_lon":28.62},
            "Kafue":                  {"min_lat":-15.88,"max_lat":-15.72,"min_lon":28.12,"max_lon":28.32},
            "Kapiri Mposhi":          {"min_lat":-13.92,"max_lat":-13.72,"min_lon":28.52,"max_lon":28.82},
            "Kasama":                 {"min_lat":-10.30,"max_lat":-10.10,"min_lon":31.10,"max_lon":31.32},
            "Kitwe":                  {"min_lat":-13.10,"max_lat":-12.72,"min_lon":28.12,"max_lon":28.52},
            "Livingstone":            {"min_lat":-17.95,"max_lat":-17.68,"min_lon":25.70,"max_lon":26.00},
            "Luanshya":               {"min_lat":-13.18,"max_lat":-13.02,"min_lon":28.32,"max_lon":28.58},
            "Lusaka (city)":          {"min_lat":-15.60,"max_lat":-15.18,"min_lon":28.10,"max_lon":28.60},
            "Mansa":                  {"min_lat":-11.32,"max_lat":-11.08,"min_lon":28.80,"max_lon":29.12},
            "Mazabuka":               {"min_lat":-15.92,"max_lat":-15.72,"min_lon":27.68,"max_lon":27.92},
            "Mongu":                  {"min_lat":-15.40,"max_lat":-15.10,"min_lon":22.95,"max_lon":23.30},
            "Mpika":                  {"min_lat":-11.92,"max_lat":-11.72,"min_lon":31.38,"max_lon":31.62},
            "Mufulira":               {"min_lat":-12.58,"max_lat":-12.38,"min_lon":28.18,"max_lon":28.42},
            "Ndola":                  {"min_lat":-13.10,"max_lat":-12.78,"min_lon":28.48,"max_lon":28.90},
            "Petauke":                {"min_lat":-14.32,"max_lat":-14.12,"min_lon":31.28,"max_lon":31.52},
            "Senanga":                {"min_lat":-16.22,"max_lat":-16.02,"min_lon":23.18,"max_lon":23.42},
            "Siavonga":               {"min_lat":-16.58,"max_lat":-16.38,"min_lon":28.58,"max_lon":28.82},
            "Solwezi":                {"min_lat":-12.32,"max_lat":-12.10,"min_lon":26.28,"max_lon":26.52},
            "Zambezi":                {"min_lat":-13.58,"max_lat":-13.38,"min_lon":23.05,"max_lon":23.28},
        }
        _area_opts = ["— Select area —"] + list(_AREA_BBOXES.keys())
        _cur_sel = st.session_state.get("_area_sel_name", "— Select area —")
        if _cur_sel not in _area_opts:
            _cur_sel = "— Select area —"
        _area_ver = st.session_state.get("_area_selector_ver", 0)
        _sel_area = st.selectbox(
            "Area", options=_area_opts,
            index=_area_opts.index(_cur_sel),
            key=f"area_selector_{_area_ver}", label_visibility="collapsed",
        )

        # Draw map — toggle button opens the map in the main content area
        _draw_open = st.session_state.get("_draw_map_open", False)
        if st.button(
            "Close Map" if _draw_open else "Open Map & Distance Tool",
            key="toggle_draw_map", use_container_width=True
        ):
            st.session_state["_draw_map_open"] = not _draw_open
            st.rerun()

        # Mini-map — Zambia outline + province centroids, highlight selected area
        if _sel_area != "— Select area —":
            st.session_state["_area_sel_name"] = _sel_area
            _btn_col, _clr_col = st.columns([3, 1])
            with _btn_col:
                _do_count = st.button("Count Features", key="count_preset_btn", use_container_width=True)
            with _clr_col:
                if st.button("✕", key="clear_area_btn", use_container_width=True):
                    st.session_state.pop("draw_bbox", None)
                    st.session_state.pop("_draw_counts", None)
                    st.session_state.pop("_draw_details", None)
                    st.session_state.pop("_area_sel_name", None)
                    # Bump version to force a fresh selectbox widget (resets to placeholder)
                    st.session_state["_area_selector_ver"] = _area_ver + 1
                    st.rerun()
            if _do_count:
                _b = {**_AREA_BBOXES[_sel_area], "measurement": f"Area: {_sel_area}"}
                st.session_state["draw_bbox"] = _b
                _bbx_str = (f"{_b['min_lon']},{_b['min_lat']},"
                            f"{_b['max_lon']},{_b['max_lat']}")
                _tok = _hub_client_module._ARCGIS_TOKEN
                _count_datasets = [
                    ("Health facilities",
                     "https://services3.arcgis.com/BU6Aadhn6tbBEdyk/arcgis/rest/services/GRID3_ZMB_HealthFac_v01beta/FeatureServer/0",
                     ["Facility_N", "Name", "FacilityNa", "facility_name", "NAME"], "Type"),
                    ("Schools",
                     "https://services3.arcgis.com/BU6Aadhn6tbBEdyk/arcgis/rest/services/GRID3_ZMB_School_v01beta/FeatureServer/0",
                     ["School_Nam", "Name", "school_name", "NAME"], "School_Typ"),
                    ("Roads",
                     "https://services3.arcgis.com/t6lYS2Pmd8iVx1fy/arcgis/rest/services/glc_ZMB_trs_roads_major_b_view/FeatureServer/0",
                     ["name", "Name", "road_name", "NAME"], "type"),
                    ("Settlements",
                     "https://services3.arcgis.com/BU6Aadhn6tbBEdyk/arcgis/rest/services/GRID3_Zambia_Operational_Settlement_Points_and_Names_Version01/FeatureServer/0",
                     ["Settlement", "name", "Name", "NAME"], "Type"),
                ]
                import requests as _req
                _counts = {}
                _details = {}
                _hdr = {
                    "Referer": "https://zmb-geowb.hub.arcgis.com",
                    "Origin":  "https://zmb-geowb.hub.arcgis.com",
                    "Accept":  "application/json",
                }
                _area_km2_draw = (
                    haversine_km(_b["min_lat"], _b["min_lon"], _b["min_lat"], _b["max_lon"]) *
                    haversine_km(_b["min_lat"], _b["min_lon"], _b["max_lat"], _b["min_lon"])
                )
                _ctr_lat_area = (_b["min_lat"] + _b["max_lat"]) / 2
                _ctr_lon_area = (_b["min_lon"] + _b["max_lon"]) / 2

                import concurrent.futures as _cf

                def _fetch_arcgis_area(args):
                    _label, _url, _name_fields, _type_field = args
                    try:
                        _p = {"geometry": _bbx_str, "geometryType": "esriGeometryEnvelope",
                              "spatialRel": "esriSpatialRelIntersects",
                              "returnCountOnly": "true", "f": "json"}
                        if _tok: _p["token"] = _tok
                        _r = _req.get(f"{_url}/query", params=_p, headers=_hdr, timeout=10)
                        _c = _r.json().get("count", 0)
                        if _c and _c > 0:
                            _fp = {"geometry": _bbx_str, "geometryType": "esriGeometryEnvelope",
                                   "spatialRel": "esriSpatialRelIntersects",
                                   "outFields": "*", "resultRecordCount": 20,
                                   "returnGeometry": "true", "f": "json"}
                            if _tok: _fp["token"] = _tok
                            _fr = _req.get(f"{_url}/query", params=_fp, headers=_hdr, timeout=12)
                            _feats = _fr.json().get("features", [])
                            _names, _subtypes, _nearest_name, _nearest_dist = [], {}, None, float("inf")
                            for _feat in _feats:
                                _props = _feat.get("attributes") or {}
                                _nm = None
                                for _nf in _name_fields:
                                    _nm = _props.get(_nf)
                                    if _nm and str(_nm).strip() not in ("None", "null", "", "0"): break
                                if _nm: _names.append(str(_nm).strip())
                                if _type_field:
                                    _st2 = _props.get(_type_field)
                                    if _st2 and str(_st2).strip() not in ("None", "null", ""):
                                        _st2 = str(_st2).strip()
                                        _subtypes[_st2] = _subtypes.get(_st2, 0) + 1
                                _geom = _feat.get("geometry") or {}
                                _fx = _geom.get("x"); _fy = _geom.get("y")
                                if _fx and _fy:
                                    _d = haversine_km(_ctr_lat_area, _ctr_lon_area, _fy, _fx)
                                    if _d < _nearest_dist:
                                        _nearest_dist = _d
                                        _nearest_name = str(_nm).strip() if _nm else None
                            return _label, _c, {"names": _names, "subtypes": _subtypes,
                                                "nearest_name": _nearest_name,
                                                "nearest_dist": _nearest_dist if _nearest_dist < float("inf") else None}
                        return _label, _c, {"names": [], "subtypes": {}, "nearest_name": None, "nearest_dist": None}
                    except Exception:
                        return _label, "—", {"names": [], "subtypes": {}, "nearest_name": None, "nearest_dist": None}

                _s2, _w2, _n2, _e2 = _b["min_lat"], _b["min_lon"], _b["max_lat"], _b["max_lon"]
                _osm_bbox2 = f"({_s2},{_w2},{_n2},{_e2})"
                _OVERPASS_MIRRORS2 = [
                    "https://overpass.kumi.systems/api/interpreter",
                    "https://overpass-api.de/api/interpreter",
                    "https://maps.mail.ru/osm/tools/overpass/api/interpreter",
                ]
                def _overpass_req2(query_str):
                    for _mirror2 in _OVERPASS_MIRRORS2:
                        try:
                            _r3 = _req.post(_mirror2, data={"data": query_str},
                                            headers={"User-Agent": "ZambiaGeoHubAI/1.0"}, timeout=20)
                            if _r3.status_code == 200: return _r3.json()
                        except Exception: continue
                    return None

                def _fetch_osm_area(label, count_q, detail_q, name_tags, default_name):
                    try:
                        _js2 = _overpass_req2(count_q)
                        if not _js2:
                            return label, "—", {"names": [], "subtypes": {}, "nearest_name": None, "nearest_dist": None}
                        _cnt2 = int((_js2.get("elements") or [{}])[0].get("tags", {}).get("total", 0))
                        if _cnt2 == 0:
                            return label, 0, {"names": [], "subtypes": {}, "nearest_name": None, "nearest_dist": None}
                        _fjs2 = _overpass_req2(detail_q)
                        _names2, _near2, _ndist2 = [], None, float("inf")
                        for _el2 in (_fjs2 or {}).get("elements", []):
                            _tags2 = _el2.get("tags") or {}
                            _nm2 = next((_tags2.get(t) for t in name_tags if _tags2.get(t)), None)
                            if _nm2: _names2.append(_nm2)
                            _lat2 = _el2.get("lat"); _lon2 = _el2.get("lon")
                            if _lat2 and _lon2:
                                _d2 = haversine_km(_ctr_lat_area, _ctr_lon_area, _lat2, _lon2)
                                if _d2 < _ndist2: _ndist2 = _d2; _near2 = _nm2 or default_name
                        return label, _cnt2, {"names": _names2, "subtypes": {}, "nearest_name": _near2,
                                              "nearest_dist": _ndist2 if _ndist2 < float("inf") else None}
                    except Exception:
                        return label, "—", {"names": [], "subtypes": {}, "nearest_name": None, "nearest_dist": None}

                _osm_tasks2 = [
                    ("Mines",
                     f'[out:json][timeout:20];(node["industrial"="mine"]{_osm_bbox2};way["industrial"="mine"]{_osm_bbox2};node["landuse"="quarry"]{_osm_bbox2};way["landuse"="quarry"]{_osm_bbox2};);out count;',
                     f'[out:json][timeout:20];(node["industrial"="mine"]{_osm_bbox2};way["industrial"="mine"]{_osm_bbox2};node["landuse"="quarry"]{_osm_bbox2};);out body 20;',
                     ["name", "operator"], "Unnamed mine"),
                    ("Dams",
                     f'[out:json][timeout:20];(node["waterway"="dam"]{_osm_bbox2};way["waterway"="dam"]{_osm_bbox2};node["man_made"="dam"]{_osm_bbox2};way["man_made"="dam"]{_osm_bbox2};);out count;',
                     f'[out:json][timeout:20];(node["waterway"="dam"]{_osm_bbox2};way["waterway"="dam"]{_osm_bbox2};);out body 20;',
                     ["name"], "Unnamed dam"),
                    ("Churches",
                     f'[out:json][timeout:20];(node["amenity"="place_of_worship"]["religion"="christian"]{_osm_bbox2};way["amenity"="place_of_worship"]["religion"="christian"]{_osm_bbox2};);out count;',
                     f'[out:json][timeout:20];(node["amenity"="place_of_worship"]["religion"="christian"]{_osm_bbox2};);out body 20;',
                     ["name"], "Unnamed church"),
                    ("Mosques",
                     f'[out:json][timeout:20];(node["amenity"="place_of_worship"]["religion"="muslim"]{_osm_bbox2};way["amenity"="place_of_worship"]["religion"="muslim"]{_osm_bbox2};);out count;',
                     f'[out:json][timeout:20];(node["amenity"="place_of_worship"]["religion"="muslim"]{_osm_bbox2};);out body 20;',
                     ["name"], "Unnamed mosque"),
                    ("Markets & shops",
                     f'[out:json][timeout:20];(node["amenity"="marketplace"]{_osm_bbox2};way["amenity"="marketplace"]{_osm_bbox2};node["shop"~"supermarket|mall|convenience|general"]{_osm_bbox2};);out count;',
                     f'[out:json][timeout:20];(node["amenity"="marketplace"]{_osm_bbox2};node["shop"~"supermarket|mall|convenience|general"]{_osm_bbox2};);out body 20;',
                     ["name"], "Unnamed market"),
                ]
                with _cf.ThreadPoolExecutor(max_workers=9) as _pool2:
                    _arcgis_futs2 = [_pool2.submit(_fetch_arcgis_area, ds) for ds in _count_datasets]
                    _osm_futs2    = [_pool2.submit(_fetch_osm_area, *t) for t in _osm_tasks2]
                    for _fut2 in _arcgis_futs2 + _osm_futs2:
                        _lbl2, _cnt2, _det2 = _fut2.result()
                        _counts[_lbl2] = _cnt2
                        _details[_lbl2] = _det2
                st.session_state["_draw_counts"]   = _counts
                st.session_state["_draw_details"]  = _details
                st.session_state["_draw_area_km2"] = _area_km2_draw
                st.session_state["_draw_centroid"] = (_ctr_lat_area, _ctr_lon_area)
                st.rerun()

    # Show area analysis results
    if st.session_state.get("_draw_counts") and st.session_state.get("_area_sel_name"):
        _dc = st.session_state["_draw_counts"]
        _dd = st.session_state.get("_draw_details", {})
        _da = st.session_state.get("_draw_area_km2", 0)
        _health_cnt = _dc.get("Health facilities", 0) if isinstance(_dc.get("Health facilities"), int) else 0
        _school_cnt = _dc.get("Schools", 0) if isinstance(_dc.get("Schools"), int) else 0
        _settle_cnt = _dc.get("Settlements", 0) if isinstance(_dc.get("Settlements"), int) else 0
        _pop_est    = _settle_cnt * 6
        _health_ratio = round(_health_cnt / _pop_est * 10000, 1) if _pop_est > 0 else None
        _school_ratio = round(_school_cnt / _settle_cnt, 2) if _settle_cnt > 0 else None
        _settle_density = (_settle_cnt / _da * 100) if _da and _da > 0 else 0
        _area_class = "Urban" if _settle_density >= 50 else ("Peri-urban" if _settle_density >= 15 else "Rural")
        st.success(f"**{st.session_state['_area_sel_name']}** — {_area_class}")
        st.caption(f"Est. population: ~{_pop_est:,}")
        st.caption(f"Health/10k people: {_health_ratio if _health_ratio is not None else 'N/A'}")
        st.caption(f"Schools/settlement: {_school_ratio if _school_ratio is not None else 'N/A'}")
        st.markdown("**Feature counts** (click to expand):")
        for _lbl3, _cnt3 in _dc.items():
            _info3 = _dd.get(_lbl3, {})
            _names3 = _info3.get("names", [])
            _subtypes3 = _info3.get("subtypes", {})
            _nearest3 = _info3.get("nearest_name")
            _ndist3 = _info3.get("nearest_dist")
            if isinstance(_cnt3, int) and _da and _da > 0:
                _exp_lbl3 = f"{_lbl3} — **{_cnt3}** · {_cnt3/_da*100:.1f}/100 km²"
            else:
                _exp_lbl3 = f"{_lbl3} — **{_cnt3}**"
            with st.expander(_exp_lbl3, expanded=False):
                if _nearest3 and _ndist3 is not None:
                    st.caption(f"📍 Nearest: {_nearest3} ({_ndist3:.1f} km)")
                if _subtypes3:
                    st.caption("🏷️ " + " · ".join(f"{k}: {v}" for k, v in sorted(_subtypes3.items(), key=lambda x: -x[1])[:4]))
                for _n3 in _names3[:10]:
                    st.caption(f"• {_n3}")
                if not _nearest3 and not _subtypes3 and not _names3:
                    st.caption("No additional details.")
        st.caption("[🔗 Zambia GeoHub](https://zmb-geowb.hub.arcgis.com)")

    st.divider()

    # New Chat button
    if st.button("New Chat", key="new_chat_btn", type="primary", use_container_width=True):
        _save_current_chat()
        st.session_state.messages = []
        st.session_state["_current_chat_id"] = str(_uuid.uuid4())
        for _k in ["draw_bbox","_draw_counts","_draw_details","_area_sel_name",
                   "uploaded_doc_text","uploaded_doc_name","uploaded_img_b64","uploaded_img_name"]:
            st.session_state.pop(_k, None)
        st.rerun()

    # Recent chats
    if st.session_state.chat_sessions:
        st.markdown('<p class="zmb-recents-label">Recents</p>', unsafe_allow_html=True)
        for _cs in st.session_state.chat_sessions[:15]:
            _is_active = _cs["id"] == st.session_state.get("_current_chat_id")
            _btn_label = ("▶ " if _is_active else "") + _cs["title"]
            if st.button(_btn_label, key=f"hist_{_cs['id']}", use_container_width=True):
                _save_current_chat()
                st.session_state.messages = list(_cs["messages"])
                st.session_state["_current_chat_id"] = _cs["id"]
                st.rerun()

    # (File upload moved to the ＋ Attach popover next to the chat input)

# ---------------------------------------------------------------------------
# Context detection — dataset passed from Hub iframe embed
#
# The Hub page embeds this app as:
#   <iframe src="https://...streamlit.app/?dataset_url=...&dataset_name=...">
#
# Supported params:
#   dataset_url  — FeatureServer layer URL of the open dataset
#   dataset_name — Human-readable name (shown in the banner)
# ---------------------------------------------------------------------------
params = st.query_params
_ctx_url  = params.get("dataset_url", "")
_ctx_name = params.get("dataset_name", "")

# Resolve the context dataset from the catalog (by URL) or create a minimal entry
context_dataset = None
if _ctx_url:
    catalog = hub.get_catalog()
    # Try exact URL match first
    for ds in catalog:
        if ds["url"].rstrip("/") == _ctx_url.rstrip("/"):
            context_dataset = ds
            break
    # If not found in catalog, create a minimal entry so we can still query it
    if context_dataset is None:
        context_dataset = {
            "id": "ctx",
            "name": _ctx_name or "Selected Dataset",
            "description": f"Dataset loaded from the Zambia GeoHub: {_ctx_url}",
            "url": _ctx_url,
            "tags": ["zambia"],
            "fields": hub._fetch_fields(_ctx_url),
            "geometry_type": "Unknown",
            "extent": {},
            "modified": "",
        }

# ---------------------------------------------------------------------------
# Session state — chat persistence via URL query param (pure Python, no JS)
# ---------------------------------------------------------------------------
# Strategy: after every message pair, serialize messages → compress → base64
# → store in st.query_params["c"]. On refresh Streamlit reads the same URL
# so the param is still there and messages are restored.
# Geojson / sample_features are stripped before encoding to keep URL short.
# ---------------------------------------------------------------------------
import json as _json_mod2
import zlib as _zlib
import base64 as _b64

_CHAT_PARAM = "c"
_MAX_CHAT_MESSAGES = 10  # keep last N messages to avoid URL length limits


def _encode_chat(messages: list) -> str:
    """Serialize messages to a compact URL-safe string."""
    slim = []
    for m in messages[-_MAX_CHAT_MESSAGES:]:
        # Exclude large fields — geojson is too big for a URL param
        entry = {k: v for k, v in m.items()
                 if k not in ("geojson", "sample_features", "docx_bytes", "pdf_bytes")
                 and isinstance(v, (str, int, float, bool, list, dict, type(None)))}
        # Truncate long AI responses to keep URL short
        if "content" in entry and isinstance(entry["content"], str) and len(entry["content"]) > 800:
            entry["content"] = entry["content"][:800] + "…"
        slim.append(entry)
    raw = _json_mod2.dumps(slim, separators=(",", ":"))
    compressed = _zlib.compress(raw.encode("utf-8"), level=6)
    return _b64.urlsafe_b64encode(compressed).decode("ascii")


def _decode_chat(encoded: str) -> list:
    """Restore messages from URL param string."""
    compressed = _b64.urlsafe_b64decode(encoded)
    raw = _zlib.decompress(compressed).decode("utf-8")
    return _json_mod2.loads(raw)


def _persist_chat():
    """Write current messages into the URL query param."""
    try:
        if st.session_state.messages:
            st.query_params[_CHAT_PARAM] = _encode_chat(st.session_state.messages)
        else:
            st.query_params.pop(_CHAT_PARAM, None)
    except Exception:
        pass


def _clear_chat_storage():
    """Remove chat from URL."""
    try:
        st.query_params.pop(_CHAT_PARAM, None)
    except Exception:
        pass


# Restore messages from URL on first load
if "messages" not in st.session_state:
    _encoded = st.query_params.get(_CHAT_PARAM, "")
    if _encoded:
        try:
            st.session_state.messages = _decode_chat(_encoded)
        except Exception:
            st.session_state.messages = []
    else:
        st.session_state.messages = []

if "edit_idx" not in st.session_state:
    st.session_state.edit_idx = None
if "stop_streaming" not in st.session_state:
    st.session_state.stop_streaming = False
if "is_generating" not in st.session_state:
    st.session_state.is_generating = False

# ---------------------------------------------------------------------------
# Fixed-bottom stop button — shown while AI is generating (matches Claude UI)
# ---------------------------------------------------------------------------
if st.session_state.get("is_generating"):
    # Dim the chat input so users know not to type while generating
    st.markdown(
        """<style>
        [data-testid="stChatInput"] textarea { opacity: 0.4; }
        [data-testid="stChatInput"] button[data-testid="stChatInputSubmitButton"] { opacity: 0; pointer-events: none; }
        </style>""",
        unsafe_allow_html=True,
    )
    # Stop button rendered in the sidebar-like area above the chat input
    _stop_top_col, _ = st.columns([1, 5])
    with _stop_top_col:
        if st.button("⏹ Stop generating", key="global_stop_btn", type="primary", use_container_width=True):
            st.session_state.stop_streaming = True
            st.session_state.is_generating = False

# ---------------------------------------------------------------------------
# Intent detection
# ---------------------------------------------------------------------------
def detect_intent(text: str, has_image: bool = False) -> str:
    t = text.lower()
    if any(w in t for w in ["report", "generate report", "write report", "create report"]):
        return "report"
    if any(w in t for w in ["summarise", "summarize", "summary", "overview", "brief"]):
        return "summary"
    if has_image:
        return "map_analysis"
    return "chat"


def _extract_comparison_locations(text: str):
    """
    Detect 'compare X and Y' or 'X vs Y' patterns.
    Returns list of location strings if found, else empty list.
    """
    t = text.lower()
    # "compare X and Y" or "X vs Y" or "X versus Y"
    for pattern in [
        _re.search(r'compare\s+([a-z][a-z\s]{2,20})\s+and\s+([a-z][a-z\s]{2,20})', t),
        _re.search(r'([a-z][a-z\s]{2,15})\s+vs\.?\s+([a-z][a-z\s]{2,15})', t),
        _re.search(r'([a-z][a-z\s]{2,15})\s+versus\s+([a-z][a-z\s]{2,15})', t),
    ]:
        if pattern:
            locs = [pattern.group(1).strip().title(), pattern.group(2).strip().title()]
            # Filter out generic words
            locs = [l for l in locs if l.lower() not in {"the", "a", "an", "all", "zambia", "africa"}]
            if len(locs) == 2:
                return locs
    return []


# ---------------------------------------------------------------------------
# Location filter helpers
# ---------------------------------------------------------------------------
import re as _re

# Zambia districts and provinces (lowercase) for map filtering
_ZAMBIA_PROVINCES = {
    "lusaka", "copperbelt", "central", "eastern", "northern", "southern",
    "western", "northwestern", "north-western", "luapula", "muchinga",
}

def _extract_location(text: str):
    """
    Extract a district/province name from a question.
    Returns (location_str, location_type) or (None, None).
    E.g. 'schools in Chadiza' → ('Chadiza', 'district')
         'hospitals in Lusaka Province' → ('Lusaka', 'province')
    Case-insensitive: works for 'kalomo' and 'Kalomo' alike.
    """
    t = text.lower()
    # Check for explicit province mention
    for prov in _ZAMBIA_PROVINCES:
        if prov in t:
            return (prov.title(), "province")
    # Look for "within X km/miles of <place>" first (most specific radius pattern)
    match_of = _re.search(r'\bof\s+([a-zA-Z][a-z]{2,}(?:\s+[a-zA-Z][a-z]{2,})?)', text, _re.IGNORECASE)
    if match_of:
        loc = _re.sub(r'\s+(?:district|province|region)\s*$', '', match_of.group(1).strip(), flags=_re.IGNORECASE).title()
        if loc.lower() not in {"zambia", "the", "all", "africa"}:
            return (loc, "district")
    # Look for "in/within/around/near/at <place>" — case-insensitive
    match = _re.search(r'\b(?:in|within|around|near|at)\s+([a-zA-Z][a-z]{2,}(?:\s+[a-zA-Z][a-z]{2,})?)', text, _re.IGNORECASE)
    if match:
        loc = _re.sub(r'\s+(?:district|province|region)\s*$', '', match.group(1).strip(), flags=_re.IGNORECASE).title()
        if loc.lower() not in {"zambia", "the", "all", "zambia province", "africa"}:
            return (loc, "district")
    # Also catch "<place> district" pattern (e.g. "Kalomo district hospitals")
    match2 = _re.search(r'\b([A-Za-z][a-z]{2,}(?:\s+[A-Za-z][a-z]{2,})?)\s+district\b', text, _re.IGNORECASE)
    if match2:
        loc = match2.group(1).strip().title()
        if loc.lower() not in {"zambia", "the", "all", "africa"}:
            return (loc, "district")
    return (None, None)


def _extract_coordinates(text: str):
    """
    Detect lat/lon coordinates typed anywhere in the question.
    Returns (lat, lon) floats or (None, None).

    Recognised patterns:
      -15.416, 28.283          (decimal, comma-separated)
      -15.416 28.283           (decimal, space-separated)
      15.4S, 28.3E             (with hemisphere letters)
      lat -15.4 lon 28.3       (labelled)
      coordinates(-15.4, 28.3) (word 'coordinates')
    Zambia bounds: lat -18 to -8, lon 21 to 34.
    """
    import re as _re_c
    t = text.strip()

    # 1. Labelled: lat/latitude ... lon/longitude
    m = _re_c.search(
        r'lat(?:itude)?\s*[:\s]\s*([-−]?\d{1,2}(?:\.\d+)?)\s*[,\s]+\s*'
        r'lon(?:gitude)?\s*[:\s]\s*([-−]?\d{1,3}(?:\.\d+)?)',
        t, _re_c.I
    )
    if m:
        lat, lon = float(m.group(1).replace('−', '-')), float(m.group(2).replace('−', '-'))
        if -18 <= lat <= -8 and 21 <= lon <= 34:
            return lat, lon

    # 2. Hemisphere letters: 15.4S, 28.3E
    m = _re_c.search(
        r'(\d{1,2}(?:\.\d+)?)\s*°?\s*([NS])\s*[,\s]+\s*(\d{1,3}(?:\.\d+)?)\s*°?\s*([EW])',
        t, _re_c.I
    )
    if m:
        lat = float(m.group(1)) * (-1 if m.group(2).upper() == 'S' else 1)
        lon = float(m.group(3)) * (-1 if m.group(4).upper() == 'W' else 1)
        if -18 <= lat <= -8 and 21 <= lon <= 34:
            return lat, lon

    # 3. Bare decimal pair — negative lat (Zambia is south of equator)
    m = _re_c.search(
        r'([-−]\d{1,2}(?:\.\d+)?)\s*[,\s]+\s*(\d{1,3}(?:\.\d+)?)',
        t
    )
    if m:
        lat, lon = float(m.group(1).replace('−', '-')), float(m.group(2))
        if -18 <= lat <= -8 and 21 <= lon <= 34:
            return lat, lon

    return None, None


def _extract_radius_km(text: str):
    """
    Detect a buffer radius in the question.
    Returns float km or None.
    Examples:
      "within 5km"  → 5.0
      "10 km radius" → 10.0
      "within 500m"  → 0.5
      "within 2 kilometres" → 2.0
    """
    import re as _re2
    t = text.lower()
    # kilometres
    m = _re2.search(r'(\d+(?:\.\d+)?)\s*(?:km|kms|kilometre|kilometres|kilometer|kilometers)', t)
    if m:
        return float(m.group(1))
    # metres → convert
    m = _re2.search(r'(\d+(?:\.\d+)?)\s*(?:m|metres|meters|metre|meter)\b', t)
    if m:
        val = float(m.group(1))
        if val >= 50:          # ignore single-digit metres that are likely typos
            return round(val / 1000, 3)
    return None


# Known coordinates for major Zambia towns/cities — used when no district polygon
# centroid is found (e.g. radius queries like "within 5km of Mongu")
_ZAMBIA_TOWN_COORDS: dict = {
    "mongu": (-15.2567, 23.1269), "lusaka": (-15.4167, 28.2833),
    "ndola": (-12.9587, 28.6366), "kitwe": (-12.8024, 28.2132),
    "kabwe": (-14.4469, 28.4464), "livingstone": (-17.8614, 25.8542),
    "chipata": (-13.6436, 32.6444), "solwezi": (-12.1731, 26.3978),
    "mansa": (-11.2, 28.8833), "kasama": (-10.2, 31.1833),
    "choma": (-16.8, 26.9833), "mazabuka": (-15.8591, 27.7497),
    "kafue": (-15.7697, 28.1861), "chingola": (-12.5279, 27.8584),
    "mufulira": (-12.5529, 28.2392), "luanshya": (-13.1353, 28.4028),
    "kalulushi": (-12.8385, 28.1013), "chililabombwe": (-12.3602, 27.8289),
    "samfya": (-11.3667, 29.55), "mwinilunga": (-11.7381, 24.4322),
    "senanga": (-16.1167, 23.2667), "sesheke": (-17.4833, 24.3),
    "zambezi": (-13.5378, 23.1069), "chavuma": (-13.0833, 22.7),
    "petauke": (-14.2478, 31.3228), "lundazi": (-12.2945, 33.175),
    "isoka": (-10.1667, 32.6333), "mpika": (-11.8309, 31.4529),
    "chinsali": (-10.5545, 32.0649), "nakonde": (-9.3667, 32.7333),
    "mbala": (-8.8333, 31.3667), "nchelenge": (-9.35, 28.7333),
    "kawambwa": (-9.7833, 29.0833), "mwense": (-10.3667, 28.7),
    "luwingu": (-10.2558, 29.9233), "kaputa": (-8.4667, 29.6833),
    "serenje": (-13.2294, 30.2364), "mkushi": (-13.6167, 29.3833),
    "kapiri mposhi": (-13.9667, 28.6833),
}


def _filter_by_location(features: list, location: str, loc_type: str) -> list:
    """Filter features to only those matching district or province."""
    loc_lower = location.lower()
    result = []
    for f in features:
        props = f.get("properties") or {}
        if loc_type == "province":
            val = (props.get("Province") or props.get("PROVINCE") or "").lower()
        else:
            val = (
                props.get("District") or props.get("DISTRICT") or
                props.get("Province") or props.get("PROVINCE") or ""
            ).lower()
        if loc_lower in val or val in loc_lower:
            result.append(f)
    return result

# ---------------------------------------------------------------------------
# Context banner — shown when a dataset is passed from the Hub page
# ---------------------------------------------------------------------------
if context_dataset:
    st.info(
        f"📍 **Context loaded:** You are viewing **{context_dataset['name']}** on the Hub. "
        f"Your questions will be answered from this dataset first.",
        icon=None,
    )

# ---------------------------------------------------------------------------
# Welcome screen — shown only on first open (no messages yet)
# ---------------------------------------------------------------------------
if not st.session_state.get("messages"):
    # Topic category buttons — open data source in new tab AND ask a question
    _TOPICS = [
        {
            "label":    "Health",
            "question": "Show me health facilities across Zambia",
            "url":      "https://zmb-geowb.hub.arcgis.com/search?q=health+facilities",
        },
        {
            "label":    "Education",
            "question": "Show me schools across Zambia",
            "url":      "https://zmb-geowb.hub.arcgis.com/search?q=schools+education",
        },
        {
            "label":    "Infrastructure",
            "question": "Show me roads and infrastructure in Zambia",
            "url":      "https://zmb-geowb.hub.arcgis.com/search?q=roads+infrastructure",
        },
        {
            "label":    "Environment",
            "question": "Show me environmental and flood risk data in Zambia",
            "url":      "https://zmb-geowb.hub.arcgis.com/search?q=environment+flood",
        },
        {
            "label":    "Reports",
            "question": "Generate a report on health facilities in Zambia",
            "url":      "https://zmb-geowb.hub.arcgis.com",
        },
    ]
    _topic_cols = st.columns(len(_TOPICS))
    for _ti, _tp in enumerate(_TOPICS):
        with _topic_cols[_ti]:
            if st.button(_tp["label"], key=f"topic_btn_{_ti}", use_container_width=True):
                st.session_state["_open_url"] = _tp["url"]
                st.session_state._pending_question = _tp["question"]
                st.rerun()

    # Hero banner
    st.markdown("""
<div class="zmb-hero">
  <h1>Zambia GeoHub AI Assistant</h1>
  <p>Ask questions about Zambia's geospatial data in plain English — get live counts, maps, and downloadable reports.</p>
</div>
""", unsafe_allow_html=True)

    # Nudge user to set an API key when none is configured (local dev only)
    _welcome_key = _resolve_ai_key(st.session_state.get("ai_provider", DEFAULT_PROVIDER))
    if not _welcome_key:
        st.info(
            "**No API key detected.** Add your API key to the `.env` file to enable AI responses. "
            "On World Bank Posit Connect the key is set in the Vars tab — nothing to do here.",
            icon="🔑",
        )

    # Compact hint — replaces the large feature cards
    st.markdown(
        '<p style="color:#6a8aaa;font-size:0.82rem;margin:10px 0 0 2px;">'
        'Try: <em>"How many health facilities are in Lusaka?"</em> · '
        '<em>"Generate a report on schools in Eastern Province"</em> · '
        'Use <strong>Compare Two Areas</strong> in the sidebar</p>',
        unsafe_allow_html=True,
    )


# ---------------------------------------------------------------------------
# Render chat history
# ---------------------------------------------------------------------------
# Index of the last assistant message — only show suggestions there
_last_assistant_idx = max(
    (i for i, m in enumerate(st.session_state.messages) if m["role"] == "assistant"),
    default=-1
)

# ---------------------------------------------------------------------------
# Draw map panel — shown in main area when toggled from sidebar
# ---------------------------------------------------------------------------
if st.session_state.get("_draw_map_open"):
    st.markdown("#### Zambia Map — Click Two Cities to Measure Distance")
    st.caption("Click a city dot to select it as **From**, click another as **To** — distance line appears instantly.")

    import pydeck as pdk
    import pandas as _pd_map

    _CITY_COORDS = {
        "Lusaka": (-15.42, 28.28), "Ndola": (-12.97, 28.64),
        "Kitwe": (-12.80, 28.21), "Livingstone": (-17.85, 25.87),
        "Chipata": (-13.64, 32.65), "Solwezi": (-12.17, 26.40),
        "Kasama": (-10.21, 31.18), "Mansa": (-11.09, 28.89),
        "Mongu": (-15.28, 23.12), "Kabwe": (-14.44, 28.45),
        "Chingola": (-12.52, 27.87), "Mufulira": (-12.55, 28.24),
        "Luanshya": (-13.13, 28.40), "Mazabuka": (-15.86, 27.76),
        "Choma": (-16.80, 26.97), "Kafue": (-15.77, 28.18),
        "Kapiri Mposhi": (-13.97, 28.69), "Mpika": (-11.90, 31.45),
        "Nakonde": (-9.33, 32.75), "Samfya": (-11.37, 29.55),
        "Serenje": (-13.23, 30.23), "Senanga": (-16.10, 23.27),
        "Siavonga": (-16.53, 28.72), "Petauke": (-14.25, 31.33),
    }

    _city_a = st.session_state.get("_map_city_a")
    _city_b = st.session_state.get("_map_city_b")

    # Color: selected cities highlighted in orange, others dark blue
    _city_df = _pd_map.DataFrame([
        {
            "name": k, "lat": v[0], "lon": v[1],
            "color": [255, 140, 0] if k in [_city_a, _city_b] else [29, 53, 87],
            "radius": 14000 if k in [_city_a, _city_b] else 9000,
        }
        for k, v in _CITY_COORDS.items()
    ])

    _layers = [
        pdk.Layer("ScatterplotLayer", data=_city_df,
                  get_position=["lon", "lat"],
                  get_radius="radius",
                  get_fill_color="color",
                  pickable=True, auto_highlight=True),
    ]

    # Draw line between selected cities
    if _city_a and _city_b and _city_a in _CITY_COORDS and _city_b in _CITY_COORDS:
        _ca, _cb = _CITY_COORDS[_city_a], _CITY_COORDS[_city_b]
        _line_df = _pd_map.DataFrame([{
            "from": [_ca[1], _ca[0]], "to": [_cb[1], _cb[0]]
        }])
        _layers.append(pdk.Layer("LineLayer", data=_line_df,
                                 get_source_position="from",
                                 get_target_position="to",
                                 get_color=[220, 50, 50],
                                 get_width=4))

    _map_event = st.pydeck_chart(
        pdk.Deck(
            layers=_layers,
            initial_view_state=pdk.ViewState(latitude=-13.5, longitude=28.5, zoom=5, pitch=0),
            map_style="https://basemaps.cartocdn.com/gl/positron-gl-style/style.json",
            tooltip={"text": "{name}"},
        ),
        use_container_width=True, height=430,
        on_select="rerun", selection_mode="single-object",
    )

    # Handle click — first click = city A, second = city B, third resets
    _clicked = (_map_event.selection.get("objects", {}).get("ScatterplotLayer") or [None])[0] if hasattr(_map_event, "selection") else None
    if _clicked:
        _clicked_name = _clicked.get("name")
        if not _city_a:
            st.session_state["_map_city_a"] = _clicked_name
            st.session_state.pop("_map_city_b", None)
        elif not _city_b and _clicked_name != _city_a:
            st.session_state["_map_city_b"] = _clicked_name
        else:
            st.session_state["_map_city_a"] = _clicked_name
            st.session_state.pop("_map_city_b", None)

    # Show result
    if _city_a and not _city_b:
        st.info(f"**{_city_a}** selected. Now click a second city.")
    elif _city_a and _city_b and _city_a in _CITY_COORDS and _city_b in _CITY_COORDS:
        _d = haversine_km(_CITY_COORDS[_city_a][0], _CITY_COORDS[_city_a][1],
                          _CITY_COORDS[_city_b][0], _CITY_COORDS[_city_b][1])
        st.success(f"**{_city_a}** → **{_city_b}**: {_d:.1f} km straight line")
        if st.button("Reset", key="reset_map_sel"):
            st.session_state.pop("_map_city_a", None)
            st.session_state.pop("_map_city_b", None)
            st.rerun()
    else:
        st.caption("Click any city dot on the map to start.")

    st.divider()

for i, msg in enumerate(st.session_state.messages):
    if msg["role"] == "user":
        with st.chat_message("user"):
            st.markdown(msg["content"])
    else:
        with st.chat_message("assistant"):
            # Intent badge
            intent = msg.get("intent", "chat")
            badge_label = {"chat": "Answer", "report": "Report", "summary": "Summary"}.get(intent, "Answer")
            badge_class = {"chat": "intent-chat", "report": "intent-report", "summary": "intent-summary"}.get(intent, "intent-chat")
            st.markdown(f'<span class="intent-badge {badge_class}">{badge_label}</span>', unsafe_allow_html=True)

            st.markdown(msg["content"])

            # Download buttons for reports
            if msg.get("docx_bytes") and msg.get("pdf_bytes"):
                st.markdown("**Download report:**")
                c1, c2 = st.columns(2)
                ds_name = msg.get("ds_name", "report")
                c1.download_button(
                    "⬇️ Word (.docx)", msg["docx_bytes"],
                    file_name=f"{ds_name.replace(' ','_')}_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"docx_{i}", use_container_width=True,
                )
                c2.download_button(
                    "⬇️ PDF", msg["pdf_bytes"],
                    file_name=f"{ds_name.replace(' ','_')}_report.pdf",
                    mime="application/pdf",
                    key=f"pdf_{i}", use_container_width=True,
                )

            # Download button for summaries
            if msg.get("summary_txt"):
                st.download_button(
                    "⬇️ Download Summary (.txt)", msg["summary_txt"],
                    file_name=f"{msg.get('ds_name','summary').replace(' ','_')}_summary.txt",
                    mime="text/plain",
                    key=f"sum_{i}",
                )

            # Data source citation
            if msg.get("data_source_url") and msg.get("ds_name"):
                st.markdown(
                    f'<div style="font-size:0.75rem;color:#888;margin-top:4px">'
                    f'📂 Source: <a href="{msg["data_source_url"]}" target="_blank">'
                    f'{msg["ds_name"]}</a>'
                    + (f' &nbsp;·&nbsp; <span style="color:{"#2d9c5c" if msg.get("data_live") else "#e07b00"}">'
                       f'{"🟢 Live data" if msg.get("data_live") else "🟡 Cached data"}</span>' )
                    + '</div>',
                    unsafe_allow_html=True,
                )

            # On-demand panel — map / table / chart buttons for chat answers
            if msg.get("intent", "chat") == "chat" and msg.get("ds_name"):
                _render_ondemand_panel(i, msg)

            # Action toolbar — SVG overlays sit in zero-height divs above each button
            st.markdown('<div class="zmb-tb"></div>', unsafe_allow_html=True)
            _prev_q_hist = next((m["content"] for m in reversed(st.session_state.messages[:i]) if m["role"] == "user"), "")
            _tool_cols = st.columns([1, 1, 1, 1, 1, 1, 1, 5])
            with _tool_cols[0]:
                st.markdown(f'<div class="zmb-ic">{_TB_ICONS[0]}</div>', unsafe_allow_html=True)
                if st.button("\u00a0", key=f"edit_{i}", help="Edit question", use_container_width=True):
                    st.session_state.edit_idx = i - 1
                    st.rerun()
            with _tool_cols[1]:
                st.markdown(f'<div class="zmb-ic">{_TB_ICONS[1]}</div>', unsafe_allow_html=True)
                if st.button("\u00a0", key=f"regen_{i}", help="Regenerate answer", use_container_width=True):
                    _prev_user = next(
                        (m["content"] for m in reversed(st.session_state.messages[:i])
                         if m["role"] == "user"), None
                    )
                    if _prev_user:
                        st.session_state.messages = st.session_state.messages[:i - 1]
                        st.session_state._pending_question = _prev_user
                        st.rerun()
            with _tool_cols[2]:
                st.markdown(f'<div class="zmb-ic">{_TB_ICONS[2]}</div>', unsafe_allow_html=True)
                _copy_text = msg.get("content", "")
                st.download_button("\u00a0", _copy_text, file_name="answer.txt",
                                   mime="text/plain", key=f"copy_{i}", help="Copy answer", use_container_width=True)
            with _tool_cols[3]:
                st.markdown(f'<div class="zmb-ic">{_TB_ICONS[3]}</div>', unsafe_allow_html=True)
                _prev_q_save = next((m["content"] for m in reversed(st.session_state.messages[:i]) if m["role"] == "user"), "Question")
                _save_text = f"Question:\n{_prev_q_save}\n\nAnswer:\n{msg.get('content','')}"
                st.download_button("\u00a0", _save_text, file_name="saved_answer.txt",
                                   mime="text/plain", key=f"save_{i}", help="Save answer", use_container_width=True)
            with _tool_cols[4]:
                st.markdown(f'<div class="zmb-ic">{_TB_ICONS[4]}</div>', unsafe_allow_html=True)
                if st.button("\u00a0", key=f"clear_{i}", help="Delete this answer", use_container_width=True):
                    _start = max(0, i - 1)
                    st.session_state.messages = st.session_state.messages[:_start] + st.session_state.messages[i + 1:]
                    st.rerun()
            _fb = msg.get("_feedback", "")
            with _tool_cols[5]:
                if _fb != "up":
                    st.markdown(f'<div class="zmb-ic">{_TB_ICONS[5]}</div>', unsafe_allow_html=True)
                    if st.button("\u00a0", key=f"fb_up_{i}", help="Good answer", use_container_width=True):
                        msg["_feedback"] = "up"
                        _fbc = st.session_state.get("_feedback_counts", {"up": 0, "down": 0})
                        _fbc["up"] += 1
                        st.session_state["_feedback_counts"] = _fbc
                        st.toast("Thanks for the feedback!")
                        st.rerun()
                else:
                    st.markdown('<div class="zmb-tb-liked">✓</div>', unsafe_allow_html=True)
            with _tool_cols[6]:
                if _fb != "down":
                    st.markdown(f'<div class="zmb-ic">{_TB_ICONS[6]}</div>', unsafe_allow_html=True)
                    if st.button("\u00a0", key=f"fb_dn_{i}", help="Bad answer", use_container_width=True):
                        msg["_feedback"] = "down"
                        _fbc = st.session_state.get("_feedback_counts", {"up": 0, "down": 0})
                        _fbc["down"] += 1
                        st.session_state["_feedback_counts"] = _fbc
                        st.toast("Thanks, we'll use this to improve.")
                        st.rerun()
                else:
                    st.markdown('<div class="zmb-tb-liked">✓</div>', unsafe_allow_html=True)

            # Row 2: follow-up suggestions — only on the most recent answer
            if i == _last_assistant_idx and msg.get("intent", "chat") == "chat" and _prev_q_hist:
                _hist_sugs = _build_suggestions(_prev_q_hist, bool(msg.get("location")), bool(msg.get("sample_features")), msg.get("ds_name", ""))
                if _hist_sugs:
                    st.markdown("<div style='margin-top:6px'></div>", unsafe_allow_html=True)
                    _spacer, _sug_col = st.columns([3, 7])
                    with _sug_col:
                        for _si, _stxt in enumerate(_hist_sugs[:3]):
                            if st.button(_stxt, key=f"hsug_{i}_{_si}", use_container_width=True):
                                st.session_state.messages.append({"role": "user", "content": _stxt})
                                st.session_state._pending_question = _stxt
                                st.rerun()

# ---------------------------------------------------------------------------
# Edit prompt UI
# ---------------------------------------------------------------------------
if st.session_state.edit_idx is not None:
    idx = st.session_state.edit_idx
    if idx < len(st.session_state.messages):
        original = st.session_state.messages[idx]["content"]
        st.markdown("---")
        st.markdown("**Edit your prompt:**")
        edited = st.text_area("", value=original, height=80, key="edit_area")
        ca, cb = st.columns([1, 6])
        with ca:
            submit = st.button("Submit", type="primary", key="submit_edit")
        with cb:
            if st.button("Cancel", key="cancel_edit"):
                st.session_state.edit_idx = None
                st.rerun()
        if submit:
            st.session_state.messages = st.session_state.messages[:idx]
            st.session_state.edit_idx = None
            st.session_state._pending_question = edited
            st.rerun()

# ---------------------------------------------------------------------------
# Process a question (either new or edited)
# ---------------------------------------------------------------------------
def process_question(question: str):
    _has_image = bool(st.session_state.get("uploaded_img_b64"))
    intent = detect_intent(question, has_image=_has_image)

    geojson = None
    map_geojson = None   # lightweight copy for map rendering (≤50 features)
    sample_features = []
    _total_count = None   # exact count from live API, set when location query succeeds
    _cross_context = {}   # flood + risk data fetched alongside main query

    # Detect location once, at the top — used in both context mode and free search mode
    _location, _loc_type = _extract_location(question)
    _comparison_locs = _extract_comparison_locations(question)  # ["Lusaka", "Copperbelt"] for compare queries
    _radius_km = _extract_radius_km(question)      # e.g. 5.0 for "within 5km"
    _buffer_center = None                           # (lat, lon) set later if radius detected

    # Coordinates typed in the question — treated like a point draw with ~5 km radius
    _coord_lat, _coord_lon = _extract_coordinates(question)
    if _coord_lat is not None and not st.session_state.get("draw_bbox"):
        _coord_r = _radius_km or 5.0   # default 5 km if no radius given
        # Build a bbox around the coordinate point
        _deg_lat = _coord_r / 111.0
        _deg_lon = _coord_r / (111.0 * abs(math.cos(math.radians(_coord_lat))) or 1)
        _coord_bbox = {
            "min_lat": _coord_lat - _deg_lat,
            "max_lat": _coord_lat + _deg_lat,
            "min_lon": _coord_lon - _deg_lon,
            "max_lon": _coord_lon + _deg_lon,
        }
        # Identify district/province from coordinates using context layers
        _coord_district, _coord_province = "", ""
        if _CONTEXT_LAYERS:
            for _cf in _CONTEXT_LAYERS[0]["geojson"].get("features", []):
                _cp = _cf.get("properties", {})
                _cg = _cf.get("geometry", {})
                if _cg.get("type") == "Polygon":
                    for _ring in _cg.get("coordinates", []):
                        if _point_in_polygon(_coord_lat, _coord_lon, _ring):
                            _coord_district = _cp.get("DISTRICT") or _cp.get("District") or ""
                            _coord_province = _cp.get("PROVINCE") or _cp.get("Province") or ""
                            break
                elif _cg.get("type") == "MultiPolygon":
                    for _poly in _cg.get("coordinates", []):
                        for _ring in _poly:
                            if _point_in_polygon(_coord_lat, _coord_lon, _ring):
                                _coord_district = _cp.get("DISTRICT") or _cp.get("District") or ""
                                _coord_province = _cp.get("PROVINCE") or _cp.get("Province") or ""
                                break
        _coord_bbox["district"] = _coord_district
        _coord_bbox["province"] = _coord_province
        # Use as draw_bbox for this question only (don't save to session state)
        if not st.session_state.get("draw_bbox"):
            st.info(
                f"📍 Detected coordinates **({_coord_lat:.4f}, {_coord_lon:.4f})** — "
                f"searching within {_coord_r:.0f} km"
                + (f" in **{_coord_district}**, {_coord_province}" if _coord_district else "") + "."
            )
        _draw_bbox = _coord_bbox
        _bbox_str_draw = (
            f"{_coord_bbox['min_lon']},{_coord_bbox['min_lat']},"
            f"{_coord_bbox['max_lon']},{_coord_bbox['max_lat']}"
        )
    else:
        _draw_bbox = st.session_state.get("draw_bbox")
    _draw_bbox_location_note = ""
    if _draw_bbox:
        if "_bbox_str_draw" not in dir():
            _bbox_str_draw = (
                f"{_draw_bbox['min_lon']},{_draw_bbox['min_lat']},"
                f"{_draw_bbox['max_lon']},{_draw_bbox['max_lat']}"
            )
        # Identify which district(s)/province the drawn area falls in
        _bbox_center_lat = (_draw_bbox['min_lat'] + _draw_bbox['max_lat']) / 2
        _bbox_center_lon = (_draw_bbox['min_lon'] + _draw_bbox['max_lon']) / 2
        _bbox_districts = []
        _bbox_provinces = []
        if _CONTEXT_LAYERS:
            for _df in _CONTEXT_LAYERS[0]["geojson"].get("features", []):
                _dp = _df.get("properties", {})
                _dd = _dp.get("DISTRICT") or _dp.get("District") or ""
                _dpr = _dp.get("PROVINCE") or _dp.get("Province") or ""
                _dg = _df.get("geometry", {})
                if _dg.get("type") == "Polygon":
                    _rings = _dg.get("coordinates", [])
                    for _ring in _rings:
                        if _point_in_polygon(_bbox_center_lat, _bbox_center_lon, _ring):
                            if _dd and _dd not in _bbox_districts:
                                _bbox_districts.append(_dd)
                            if _dpr and _dpr not in _bbox_provinces:
                                _bbox_provinces.append(_dpr)
                            break
                elif _dg.get("type") == "MultiPolygon":
                    for _poly in _dg.get("coordinates", []):
                        for _ring in _poly:
                            if _point_in_polygon(_bbox_center_lat, _bbox_center_lon, _ring):
                                if _dd and _dd not in _bbox_districts:
                                    _bbox_districts.append(_dd)
                                if _dpr and _dpr not in _bbox_provinces:
                                    _bbox_provinces.append(_dpr)
                                break

        if _bbox_districts:
            _draw_bbox_location_note = (
                f"District(s): {', '.join(_bbox_districts)}"
                + (f" | Province(s): {', '.join(_bbox_provinces)}" if _bbox_provinces else "")
            )
            st.info(
                f"Querying within drawn area — **{_draw_bbox_location_note}** "
                f"({_draw_bbox['min_lat']:.3f}–{_draw_bbox['max_lat']:.3f}°N, "
                f"{_draw_bbox['min_lon']:.3f}–{_draw_bbox['max_lon']:.3f}°E)"
            )
        else:
            st.info(
                f"Querying within drawn area: "
                f"lat {_draw_bbox['min_lat']:.3f}–{_draw_bbox['max_lat']:.3f}, "
                f"lon {_draw_bbox['min_lon']:.3f}–{_draw_bbox['max_lon']:.3f}"
            )

    if context_dataset:
        # Hub context mode — direct request so no hub/client.py reload needed
        datasets = [context_dataset]
        with st.spinner(f"Loading data from '{context_dataset['name']}'..."):
            try:
                import requests as _req
                _base_url = context_dataset["url"].rstrip("/")
                if _base_url.endswith("/query"):
                    _base_url = _base_url[:-6]
                if _location and _loc_type == "district":
                    _where = f"District='{_location}' OR DISTRICT='{_location}'"
                elif _location and _loc_type == "province":
                    _where = f"Province='{_location}' OR PROVINCE='{_location}'"
                else:
                    _where = "1=1"
                _ctx_tok = _hub_client_module._ARCGIS_TOKEN
                _ctx_tok_p = {"token": _ctx_tok} if _ctx_tok and any(
                    org in _base_url for org in ("iQ1dY19aHwbSDYIF", "P3ePLMYs2RVChkJx")
                ) else {}
                _resp = _req.get(
                    f"{_base_url}/query",
                    params={"where": _where, "outFields": "*",
                            "resultRecordCount": 200, "f": "geojson", **_ctx_tok_p},
                    headers={"Referer": "https://zmb-geowb.hub.arcgis.com",
                             "Origin": "https://zmb-geowb.hub.arcgis.com"},
                    timeout=30,
                )
                _resp.raise_for_status()
                geojson = _resp.json()
                live_feats = geojson.get("features", [])
                if live_feats and _location:
                    st.info(f"🌐 Showing {len(live_feats)} live records for {_location}.")
                sample_features = geojson_to_sample_rows(geojson, n=len(live_feats))
                map_geojson = {"type": "FeatureCollection", "features": live_feats[:200]}
            except Exception as _ctx_e:
                st.warning(f"⚠️ Could not load data for {_location or 'this dataset'}: {_ctx_e}")
    else:
        # Free search mode — find the most relevant dataset
        with st.spinner("Searching Zambia GeoHub..."):
            try:
                datasets = hub.search_datasets(question, max_results=5)
            except Exception:
                datasets = []

        from hub.client import _load_static, _POI_TYPE_MAP_MODULE, _SUBJECT_BOOST_MODULE, _SEED_CATALOG

        _poi_type = ""
        for kw, ptype in _POI_TYPE_MAP_MODULE.items():
            if kw in question.lower():
                _poi_type = ptype
                break

        def _find_static(q_lower):
            # Use _SEED_CATALOG for subject-boost lookup — it has correct layer numbers.
            # The dynamic catalog from hub.get_catalog() enumerates all layers of each
            # FeatureServer service (0, 1, 2, …) and layer 0 always appears first,
            # which is wrong for multi-layer services like mines (/12), poverty (/50), etc.
            catalog = hub.get_catalog()
            for kw, frag in _SUBJECT_BOOST_MODULE.items():
                if kw in q_lower:
                    for ds in _SEED_CATALOG:
                        if frag in ds["url"]:
                            sd = _load_static(ds["url"], poi_type=_poi_type)
                            if sd and sd.get("features"):
                                return sd, ds
                            else:
                                # No local static file but this IS the right dataset —
                                # return it as the candidate so live fetch runs on it.
                                return None, ds
            for kw in _POI_TYPE_MAP_MODULE:
                if kw in q_lower:
                    for ds in catalog:
                        if "Points_of_Interest" in ds["url"]:
                            sd = _load_static(ds["url"], poi_type=_poi_type)
                            if sd and sd.get("features"):
                                return sd, ds
                    break
            for candidate in datasets:
                sd = _load_static(candidate["url"], poi_type=_poi_type)
                if sd and sd.get("features"):
                    return sd, candidate
            return None, None

        # If location is mentioned, always try live fetch first with a district/province
        # filter — this gives accurate counts for ALL 116 Zambia districts regardless of
        # whether the local static sample covers that area.
        # Fallback chain: live filtered → static filtered → static full
        # EXCEPTION: when a draw_bbox is active, skip the location path — the bbox
        # spatial filter is more accurate and handles the static fallback + radius filter.
        if _location and not _draw_bbox:
            _static_data, _static_candidate = _find_static(question.lower())
            _live_candidate = _static_candidate or (datasets[0] if datasets else None)

            # 1. Direct location-filtered fetch — handles both district-field datasets
            # (schools, health, POI) and spatial datasets (settlements, which have no
            # District field and need a bounding-box query against the district polygon).
            _live_error = ""
            if _live_candidate:
                with st.spinner(f"Querying live data for {_location}..."):
                    try:
                        import requests as _req
                        _base_url = _live_candidate["url"].rstrip("/")
                        if _base_url.endswith("/query"):
                            _base_url = _base_url[:-6]
                        _headers = {"Referer": "https://zmb-geowb.hub.arcgis.com",
                                    "Origin": "https://zmb-geowb.hub.arcgis.com"}

                        # Initialise so we never hit NameError if a branch doesn't set these
                        live_feats = []
                        _gjson = {}

                        # Datasets with no District/Province field use a bounding-box
                        # spatial query against the district polygon from districts.json.
                        # This covers: Settlements, Building Footprints, Population,
                        # DRE Atlas, Wealth Index, and other raster/polygon datasets.
                        _BBOX_DATASETS = (
                            "Settlement", "Building_Footprints", "Population",
                            "zambia_dre", "Relative_Wealth", "Aquifer",
                            "Microgrids", "zmb_dams", "zmb_mines",
                        )
                        _needs_bbox = any(kw in _live_candidate.get("url", "") for kw in _BBOX_DATASETS)
                        # Token for private-org datasets
                        _tok = _hub_client_module._ARCGIS_TOKEN
                        _tok_params = {"token": _tok} if _tok and any(
                            org in _base_url for org in ("iQ1dY19aHwbSDYIF", "P3ePLMYs2RVChkJx")
                        ) else {}

                        if _needs_bbox and _loc_type in ("district", "province"):
                            # Find the matching boundary polygon(s) and compute bounding box
                            _boundary_feats = _CONTEXT_LAYERS[0]["geojson"]["features"] if _CONTEXT_LAYERS else []
                            _field = "DISTRICT" if _loc_type == "district" else "PROVINCE"
                            _alt_field = "District" if _loc_type == "district" else "Province"
                            # For provinces, collect ALL matching district features to get full extent
                            _matched_feats = [
                                f for f in _boundary_feats
                                if _location.lower() in (
                                    f["properties"].get(_field) or
                                    f["properties"].get(_alt_field) or ""
                                ).lower()
                            ]
                            from utils.geo_utils import _polygon_bounds
                            # Merge bounds across all matched features
                            _bnds = None
                            for _mf in _matched_feats:
                                _b = _polygon_bounds(_mf["geometry"])
                                if _b:
                                    if _bnds is None:
                                        _bnds = [list(_b[0]), list(_b[1])]
                                    else:
                                        _bnds[0][0] = min(_bnds[0][0], _b[0][0])
                                        _bnds[0][1] = min(_bnds[0][1], _b[0][1])
                                        _bnds[1][0] = max(_bnds[1][0], _b[1][0])
                                        _bnds[1][1] = max(_bnds[1][1], _b[1][1])
                            if _bnds:
                                _bbox_str = f"{_bnds[0][1]},{_bnds[0][0]},{_bnds[1][1]},{_bnds[1][0]}"
                                _bbox_params = {
                                    "geometry": _bbox_str,
                                    "geometryType": "esriGeometryEnvelope",
                                    "spatialRel": "esriSpatialRelIntersects",
                                    **_tok_params,
                                }
                                try:
                                    _cnt_resp = _req.get(f"{_base_url}/query",
                                        params={**_bbox_params, "returnCountOnly": "true", "f": "json"},
                                        headers=_headers, timeout=15)
                                    _cnt_data = _cnt_resp.json()
                                    if "count" in _cnt_data:
                                        _total_count = _cnt_data["count"]
                                except Exception:
                                    pass
                                _resp = _req.get(f"{_base_url}/query",
                                    params={**_bbox_params, "outFields": "*",
                                            "resultRecordCount": 30, "f": "geojson"},
                                    headers=_headers, timeout=30)
                                _resp.raise_for_status()
                                _gjson = _resp.json()
                                if "error" in _gjson:
                                    _err = _gjson["error"]
                                    raise RuntimeError(f"API error {_err.get('code')}: {_err.get('message')}")
                                live_feats = _gjson.get("features", [])
                        else:
                            # Standard district/province field filter
                            _where = (
                                f"District='{_location}' OR DISTRICT='{_location}'"
                                if _loc_type == "district"
                                else f"Province='{_location}' OR PROVINCE='{_location}'"
                            )
                            # Count-only query for exact total
                            try:
                                _cnt_resp = _req.get(f"{_base_url}/query",
                                    params={"where": _where, "returnCountOnly": "true", "f": "json",
                                            **_tok_params},
                                    headers=_headers, timeout=15)
                                _cnt_data = _cnt_resp.json()
                                if "count" in _cnt_data:
                                    _total_count = _cnt_data["count"]
                            except Exception:
                                pass
                            # Sample fetch — up to 200 records for table + AI analysis
                            _resp = _req.get(f"{_base_url}/query",
                                params={"where": _where, "outFields": "*",
                                        "resultRecordCount": 200, "f": "geojson",
                                        **_tok_params},
                                headers=_headers, timeout=30)
                            _resp.raise_for_status()
                            _gjson = _resp.json()
                            if "error" in _gjson:
                                _err = _gjson["error"]
                                raise RuntimeError(f"API error {_err.get('code')}: {_err.get('message')}")
                            live_feats = _gjson.get("features", [])

                        if live_feats and "error" not in _gjson:
                            geojson = _gjson
                            st.session_state["_last_fetch_was_live"] = True
                            # Spatially assign district/province to features that lack them
                            if _CONTEXT_LAYERS:
                                assign_districts(live_feats, _CONTEXT_LAYERS[0]["geojson"])
                            sample_features = geojson_to_sample_rows(geojson, n=len(live_feats))
                            map_geojson = {"type": "FeatureCollection", "features": live_feats[:200]}
                            datasets = [_live_candidate] + [d for d in datasets if d != _live_candidate]
                            _count_label = f" (total in full dataset: {_total_count:,})" if _total_count is not None else ""
                            st.info(f"🌐 Showing {len(live_feats)} live records for {_location}{_count_label}.")
                    except Exception as _e:
                        _live_error = str(_e)

            # 1b. Cross-dataset context — flood and risk fetched alongside main dataset.
            # Flood DistName field is stored in ALL CAPS so we use UPPER() for matching.
            # If the specific district isn't flood-prone, we fetch its province neighbours
            # so the AI can say "Mongu isn't listed but 4 neighbouring districts are."
            _cross_context = {}
            if _location:
                try:
                    import requests as _req
                    _headers = {"Referer": "https://zmb-geowb.hub.arcgis.com",
                                "Origin": "https://zmb-geowb.hub.arcgis.com"}
                    catalog = hub.get_catalog()

                    # Flood-prone districts — DistName stored as UPPER CASE
                    _flood_ds = next((d for d in catalog if "Flood" in d["url"]), None)
                    if _flood_ds:
                        _flood_base = _flood_ds["url"].rstrip("/")
                        # Try exact district first (case-insensitive via UPPER())
                        if _loc_type == "district":
                            _flood_where = f"UPPER(DistName)='{_location.upper()}'"
                        else:
                            _flood_where = f"UPPER(PovName)='{_location.upper()}'"
                        _fr = _req.get(f"{_flood_base}/query",
                            params={"where": _flood_where, "outFields": "DistName,PovName",
                                    "resultRecordCount": 20, "f": "json"},
                            headers=_headers, timeout=10)
                        _fd = _fr.json()
                        if _fd.get("features"):
                            _cross_context["flood"] = [f["attributes"] for f in _fd["features"]]
                            _cross_context["flood_note"] = f"{_location} IS listed as flood-prone."
                        elif _loc_type == "district":
                            # District not flood-prone — fetch province neighbours for context
                            _prov = next(
                                (f["properties"].get("PROVINCE") for f in _CONTEXT_LAYERS[0]["geojson"]["features"]
                                 if _location.lower() in (f["properties"].get("DISTRICT") or "").lower()),
                                None
                            )
                            if _prov:
                                _prov_where = f"UPPER(PovName)='{_prov.upper()}'"
                                _fr2 = _req.get(f"{_flood_base}/query",
                                    params={"where": _prov_where, "outFields": "DistName,PovName",
                                            "resultRecordCount": 20, "f": "json"},
                                    headers=_headers, timeout=10)
                                _fd2 = _fr2.json()
                                if _fd2.get("features"):
                                    _cross_context["flood"] = [f["attributes"] for f in _fd2["features"]]
                                    _cross_context["flood_note"] = (
                                        f"{_location} district itself is NOT listed as flood-prone in the dataset. "
                                        f"However, these districts in {_prov} Province are: "
                                        f"{', '.join(f['attributes'].get('DistName','') for f in _fd2['features'])}."
                                    )
                                else:
                                    _cross_context["flood_note"] = (
                                        f"Neither {_location} nor any district in {_prov} Province "
                                        f"appears in the flood-prone districts dataset."
                                    )

                    # Risk layers — province-level aggregated data
                    _risk_ds = next((d for d in catalog if "Risk" in d["url"] and "Lusaka" not in d["url"]), None)
                    if _risk_ds:
                        _prov_for_risk = _location if _loc_type == "province" else next(
                            (f["properties"].get("PROVINCE") for f in _CONTEXT_LAYERS[0]["geojson"]["features"]
                             if _location.lower() in (f["properties"].get("DISTRICT") or "").lower()),
                            None
                        )
                        _risk_where = f"UPPER(PovName)='{_prov_for_risk.upper()}'" if _prov_for_risk else "1=1"
                        _rr = _req.get(f"{_risk_ds['url'].rstrip('/')}/query",
                            params={"where": _risk_where, "outFields": "*",
                                    "resultRecordCount": 5, "f": "json"},
                            headers=_headers, timeout=10)
                        _rd = _rr.json()
                        if _rd.get("features"):
                            _cross_context["risk"] = [f["attributes"] for f in _rd["features"]]

                    # Settlement fetch — works for both districts and provinces.
                    # When flood data is present, fetches settlements for flood-prone
                    # districts specifically so the AI can name actual communities at risk.
                    _settle_ds = next((d for d in catalog if "Settlement_Points" in d["url"] or
                                      ("Settlement" in d["url"] and "Extent" not in d["url"])), None)
                    if _settle_ds:
                        _settle_base = _settle_ds["url"].rstrip("/")
                        from utils.geo_utils import _polygon_bounds

                        # Build list of districts to fetch settlements for:
                        # — if flood-prone districts found, use those (most relevant)
                        # — otherwise use the asked district/province directly
                        _target_districts = []
                        _flood_records = _cross_context.get("flood", [])
                        if _flood_records and _loc_type == "province":
                            # Flood-prone districts in this province — fetch each
                            _target_districts = [
                                r.get("DistName", "").title()
                                for r in _flood_records if r.get("DistName")
                            ]
                        elif _loc_type == "district":
                            _target_districts = [_location]

                        _all_settle_feats = []
                        _total_settle_count = 0
                        for _td in _target_districts:
                            _td_feat = next(
                                (f for f in _CONTEXT_LAYERS[0]["geojson"]["features"]
                                 if _td.lower() in (f["properties"].get("DISTRICT") or "").lower()),
                                None
                            )
                            if not _td_feat:
                                continue
                            _bnds = _polygon_bounds(_td_feat["geometry"])
                            if not _bnds:
                                continue
                            _sbbox = f"{_bnds[0][1]},{_bnds[0][0]},{_bnds[1][1]},{_bnds[1][0]}"
                            # Count
                            try:
                                _scnt = _req.get(f"{_settle_base}/query",
                                    params={"geometry": _sbbox, "geometryType": "esriGeometryEnvelope",
                                            "spatialRel": "esriSpatialRelContains",
                                            "returnCountOnly": "true", "f": "json"},
                                    headers=_headers, timeout=15)
                                _c = _scnt.json().get("count", 0)
                                _total_settle_count += _c
                                _cross_context.setdefault("settlement_counts_by_district", {})[_td] = _c
                            except Exception:
                                pass
                            # Sample (up to 15 per district, cap total at 50)
                            if len(_all_settle_feats) < 50:
                                try:
                                    _sresp = _req.get(f"{_settle_base}/query",
                                        params={"geometry": _sbbox, "geometryType": "esriGeometryEnvelope",
                                                "spatialRel": "esriSpatialRelContains",
                                                "outFields": "*", "resultRecordCount": 15, "f": "geojson"},
                                        headers=_headers, timeout=20)
                                    _all_settle_feats.extend(_sresp.json().get("features", []))
                                except Exception:
                                    pass

                        # Province-level total (when no specific districts targeted)
                        if not _target_districts and _loc_type == "province":
                            try:
                                _sw = f"Province='{_location}' OR PROVINCE='{_location}'"
                                _scnt2 = _req.get(f"{_settle_base}/query",
                                    params={"where": _sw, "returnCountOnly": "true", "f": "json"},
                                    headers=_headers, timeout=15)
                                _total_settle_count = _scnt2.json().get("count", 0)
                                # Sample
                                _sr2 = _req.get(f"{_settle_base}/query",
                                    params={"where": _sw, "outFields": "*",
                                            "resultRecordCount": 30, "f": "geojson"},
                                    headers=_headers, timeout=20)
                                _all_settle_feats = _sr2.json().get("features", [])
                            except Exception:
                                pass

                        if _total_settle_count:
                            _cross_context["settlement_count"] = _total_settle_count
                        if _all_settle_feats:
                            _cross_context["settlement_sample"] = [
                                f["properties"] for f in _all_settle_feats[:8]
                            ]
                            _cross_context["settlement_geojson"] = {
                                "type": "FeatureCollection", "features": _all_settle_feats[:50]
                            }

                    # Roads fetch — spatial bbox query (no District field), LineString geometry.
                    # Fetches road names, numbers, surface, and functional class for the queried
                    # district or province.
                    _road_base = "https://services3.arcgis.com/t6lYS2Pmd8iVx1fy/arcgis/rest/services/glc_ZMB_trs_roads_major_b_view/FeatureServer/0"
                    _road_target_districts = (
                        [_location] if _loc_type == "district"
                        else [r.get("DistName", "").title() for r in _cross_context.get("flood", []) if r.get("DistName")]
                        if _loc_type == "province" and _cross_context.get("flood")
                        else []
                    )
                    _all_road_feats = []
                    _total_road_count = 0
                    try:
                        if _road_target_districts:
                            for _rd_dist in _road_target_districts:
                                _rd_feat = next(
                                    (f for f in _CONTEXT_LAYERS[0]["geojson"]["features"]
                                     if _rd_dist.lower() in (f["properties"].get("DISTRICT") or "").lower()),
                                    None
                                )
                                if not _rd_feat:
                                    continue
                                _rd_bnds = _polygon_bounds(_rd_feat["geometry"])
                                if not _rd_bnds:
                                    continue
                                _rd_bbox = f"{_rd_bnds[0][1]},{_rd_bnds[0][0]},{_rd_bnds[1][1]},{_rd_bnds[1][0]}"
                                # Count
                                _rcnt = _req.get(f"{_road_base}/query",
                                    params={"geometry": _rd_bbox, "geometryType": "esriGeometryEnvelope",
                                            "spatialRel": "esriSpatialRelIntersects",
                                            "returnCountOnly": "true", "f": "json"},
                                    headers=_headers, timeout=15)
                                _total_road_count += _rcnt.json().get("count", 0)
                                # Sample (up to 20 roads per district)
                                if len(_all_road_feats) < 60:
                                    _rresp = _req.get(f"{_road_base}/query",
                                        params={"geometry": _rd_bbox, "geometryType": "esriGeometryEnvelope",
                                                "spatialRel": "esriSpatialRelIntersects",
                                                "outFields": "name,roadnoloc,roadno,fclass,surface,numlanes,speedlimitkmh",
                                                "resultRecordCount": 20, "f": "geojson"},
                                        headers=_headers, timeout=20)
                                    _all_road_feats.extend(_rresp.json().get("features", []))
                        elif _loc_type == "province":
                            # Province-level: use a province bbox from context layers
                            # Find any district in province to get rough bbox
                            _prov_feats = [
                                f for f in _CONTEXT_LAYERS[0]["geojson"]["features"]
                                if _location.lower() in (f["properties"].get("PROVINCE") or "").lower()
                            ]
                            if _prov_feats:
                                # Merge all district bounds into province bbox
                                all_lats, all_lons = [], []
                                for _pf in _prov_feats:
                                    _pb = _polygon_bounds(_pf["geometry"])
                                    if _pb:
                                        all_lats += [_pb[0][0], _pb[1][0]]
                                        all_lons += [_pb[0][1], _pb[1][1]]
                                if all_lats:
                                    _prov_bbox = f"{min(all_lons)},{min(all_lats)},{max(all_lons)},{max(all_lats)}"
                                    _rcnt2 = _req.get(f"{_road_base}/query",
                                        params={"geometry": _prov_bbox, "geometryType": "esriGeometryEnvelope",
                                                "spatialRel": "esriSpatialRelIntersects",
                                                "returnCountOnly": "true", "f": "json"},
                                        headers=_headers, timeout=15)
                                    _total_road_count = _rcnt2.json().get("count", 0)
                                    _rresp2 = _req.get(f"{_road_base}/query",
                                        params={"geometry": _prov_bbox, "geometryType": "esriGeometryEnvelope",
                                                "spatialRel": "esriSpatialRelIntersects",
                                                "outFields": "name,roadnoloc,roadno,fclass,surface,numlanes,speedlimitkmh",
                                                "resultRecordCount": 40, "f": "geojson"},
                                        headers=_headers, timeout=20)
                                    _all_road_feats = _rresp2.json().get("features", [])
                    except Exception:
                        pass

                    if _total_road_count:
                        _cross_context["road_count"] = _total_road_count
                    if _all_road_feats:
                        _cross_context["road_sample"] = [
                            {k: v for k, v in (f.get("properties") or {}).items() if v is not None}
                            for f in _all_road_feats[:10]
                        ]
                        _cross_context["road_geojson"] = {
                            "type": "FeatureCollection", "features": _all_road_feats[:60]
                        }
                except Exception:
                    pass

            # 2. Live failed or was blocked — filter static data by location
            if not sample_features and _static_data and _static_data.get("features"):
                loc_feats = _filter_by_location(_static_data["features"], _location, _loc_type)
                if loc_feats:
                    sample_features = geojson_to_sample_rows(
                        {"type": "FeatureCollection", "features": loc_feats}, n=len(loc_feats)
                    )
                    map_geojson = {"type": "FeatureCollection", "features": loc_feats[:50]}
                    if _static_candidate:
                        datasets = [_static_candidate] + [d for d in datasets if d != _static_candidate]
                    st.session_state["_last_fetch_was_live"] = False
                    st.info(f"📦 Showing {len(loc_feats)} pre-loaded records for {_location} (live server unavailable).")

            # 3. Location not found anywhere
            if not sample_features:
                if _live_candidate or _static_candidate:
                    datasets = [_live_candidate or _static_candidate] + [
                        d for d in datasets if d not in (_live_candidate, _static_candidate)
                    ]
                _err_detail = f" (error: {_live_error})" if _live_error else ""
                st.warning(
                    f"⚠️ Could not load live data for **{_location}**{_err_detail}. "
                    f"The live GeoHub server may be temporarily unavailable."
                )
                if _radius_km:
                    # For radius queries: don't set a placeholder — let the countrywide
                    # static fallback (below) load all features, then the haversine filter
                    # will find what's within range of the named location's coordinates.
                    pass
                else:
                    # Non-radius query: tell Claude data is unavailable for this location
                    sample_features = [{"_note": (
                        f"No data could be retrieved for {_location}. "
                        f"The live GeoHub server is temporarily unavailable and the pre-loaded "
                        f"offline data does not cover {_location}. "
                        f"Tell the user: the live server is down, no offline data exists for {_location}, "
                        f"and they should try again later. Do NOT show data from other provinces or districts."
                    )}]

        # No location: try live fetch, fall back to static
        if not sample_features:
            _fetch_errors = []
            for candidate in datasets:
                with st.spinner(f"Loading data from '{candidate['name']}'..."):
                    try:
                        if _draw_bbox:
                            # Use drawn bounding box as spatial filter
                            import requests as _req
                            _base_url = candidate["url"].rstrip("/")
                            _query_url = f"{_base_url}/query"
                            _tok = _hub_client_module._ARCGIS_TOKEN
                            _draw_params = {
                                "geometry": _bbox_str_draw,
                                "geometryType": "esriGeometryEnvelope",
                                "spatialRel": "esriSpatialRelIntersects",
                                "outFields": "*",
                                "resultRecordCount": 200,
                                "f": "geojson",
                            }
                            if _tok:
                                _draw_params["token"] = _tok
                            _draw_headers = {"Referer": "https://zmb-geowb.hub.arcgis.com",
                                             "User-Agent": "Mozilla/5.0"}
                            _draw_resp = _req.get(_query_url, params=_draw_params,
                                                  headers=_draw_headers, timeout=30)
                            _draw_resp.raise_for_status()
                            geojson = _draw_resp.json()
                            if "error" in geojson:
                                _err = geojson["error"]
                                raise RuntimeError(f"API error {_err.get('code')}: {_err.get('message')}")
                        else:
                            geojson = hub.fetch_geojson(candidate["url"], query_hint=question)
                        sample_features = geojson_to_sample_rows(geojson, n=200)
                        if sample_features:
                            map_geojson = {"type": "FeatureCollection", "features": geojson.get("features", [])[:50]}
                            datasets = [candidate] + [d for d in datasets if d != candidate]
                            break
                    except Exception as e:
                        _fetch_errors.append(f"{candidate['name']}: {e}")

        # Last resort: static fallback
        if not sample_features:
            st.session_state["_last_fetch_was_live"] = False
            _static_data, _static_candidate = _find_static(question.lower())
            if _static_data and _static_data.get("features"):
                _static_feats = _static_data["features"]
                # Filter static features by drawn bbox if one is active
                if _draw_bbox:
                    def _feat_in_bbox(feat, bb):
                        _g = feat.get("geometry") or {}
                        _gc = _g.get("coordinates", [])
                        _gt = _g.get("type", "")
                        if _gt == "Point" and _gc:
                            lon, lat = _gc[0], _gc[1]
                            return (bb["min_lon"] <= lon <= bb["max_lon"] and
                                    bb["min_lat"] <= lat <= bb["max_lat"])
                        return False
                    _bbox_feats = [f for f in _static_feats if _feat_in_bbox(f, _draw_bbox)]
                    if _bbox_feats:
                        _static_feats = _bbox_feats
                        st.info(f"📦 Found **{len(_bbox_feats)}** records in your drawn area from local data (live server temporarily unavailable).")
                    else:
                        # Bbox coordinate filter found nothing — try filtering by district/province name
                        _area_district = _draw_bbox.get("district", "")
                        _area_province = _draw_bbox.get("province", "")
                        _district_feats = []
                        if _area_district:
                            _district_feats = [
                                f for f in _static_feats
                                if _area_district.lower() in str(f.get("properties", {}).get("District", "")).lower()
                                or _area_district.lower() in str(f.get("properties", {}).get("DISTRICT", "")).lower()
                            ]
                        if not _district_feats and _area_province:
                            _district_feats = [
                                f for f in _static_feats
                                if _area_province.lower() in str(f.get("properties", {}).get("Province", "")).lower()
                                or _area_province.lower() in str(f.get("properties", {}).get("PROVINCE", "")).lower()
                            ]
                        if _district_feats:
                            _static_feats = _district_feats
                            _label = _area_district or _area_province
                            st.info(f"📦 Found **{len(_district_feats)}** records in **{_label}** from local data (live server temporarily unavailable).")
                        else:
                            st.info("📦 No local records found in your drawn area — showing countrywide sample data.")
                else:
                    st.info("📦 Using pre-loaded sample data (live server temporarily unavailable).")
                sample_features = geojson_to_sample_rows(
                    {"type": "FeatureCollection", "features": _static_feats},
                    n=len(_static_feats)
                )
                # Use full static set for map_geojson so the radius filter
                # (applied below) can search all features, not just the first 50
                map_geojson = {"type": "FeatureCollection", "features": _static_feats}
                if _static_candidate:
                    datasets = [_static_candidate] + [d for d in datasets if d != _static_candidate]

    # ---- Buffer / proximity filter ----
    # If the question contains a radius (e.g. "within 5km of Mongu"), compute the
    # buffer center from the highlighted district centroid or the first fetched point,
    # filter map_geojson to only features inside the radius, and annotate each with
    # its distance so Claude can mention exact distances in its answer.
    if _radius_km and map_geojson and map_geojson.get("features"):
        # Determine center: prefer the district polygon centroid (most accurate for
        # location-named queries like "schools within 10km of Kalomo").
        _center = None
        if _location and _CONTEXT_LAYERS:
            _dist_feat = next(
                (f for f in _CONTEXT_LAYERS[0]["geojson"]["features"]
                 if _location.lower() in (
                     (f["properties"].get("DISTRICT") or f["properties"].get("PROVINCE") or "")
                 ).lower()),
                None,
            )
            if _dist_feat and _dist_feat.get("geometry"):
                _center = polygon_centroid(_dist_feat["geometry"])

        # Fallback: centroid of drawn bbox (covers "within the buffer" phrasing)
        if not _center and _draw_bbox:
            _center = (
                (_draw_bbox["min_lat"] + _draw_bbox["max_lat"]) / 2,
                (_draw_bbox["min_lon"] + _draw_bbox["max_lon"]) / 2,
            )

        # Fallback: known town/city coordinate lookup
        if not _center and _location:
            _town_coord = _ZAMBIA_TOWN_COORDS.get(_location.lower())
            if _town_coord:
                _center = _town_coord
                st.info(f"📍 Using known coordinates for **{_location}** ({_town_coord[0]:.4f}, {_town_coord[1]:.4f}) to apply {_radius_km} km radius filter.")

        # Fallback: centroid of all fetched point features
        if not _center:
            _pt_feats = [
                f for f in map_geojson["features"]
                if (f.get("geometry") or {}).get("type") == "Point"
            ]
            if _pt_feats:
                _lats = [f["geometry"]["coordinates"][1] for f in _pt_feats]
                _lons = [f["geometry"]["coordinates"][0] for f in _pt_feats]
                _center = (sum(_lats) / len(_lats), sum(_lons) / len(_lons))

        if _center:
            _buffer_center = _center
            _within = features_within_km(map_geojson["features"], _center[0], _center[1], _radius_km)

            if _within:
                # Annotate each feature with distance_km and coordinates
                for feat, dist_km in _within:
                    if feat.get("properties") is not None:
                        feat["properties"]["distance_km"] = dist_km
                        # Add lat/lon from geometry so coordinates appear in table/popup
                        geom = feat.get("geometry") or {}
                        if geom.get("type") == "Point":
                            coords = geom.get("coordinates", [])
                            if len(coords) >= 2:
                                feat["properties"]["longitude"] = round(coords[0], 6)
                                feat["properties"]["latitude"] = round(coords[1], 6)
                map_geojson = {"type": "FeatureCollection", "features": [f for f, _ in _within]}
                # Rebuild sample_features from the filtered + annotated set
                sample_features = geojson_to_sample_rows(map_geojson, n=len(_within))
                st.info(
                    f"📍 Buffer filter: {len(_within)} feature(s) found within "
                    f"**{_radius_km} km** of {_location or 'selected point'}."
                )
            else:
                st.info(
                    f"📍 No features found within {_radius_km} km of "
                    f"{_location or 'selected point'} in the offline data "
                    f"(live server unavailable — try again later)."
                )
                # Clear sample_features so Claude doesn't receive data from the wrong
                # location (e.g. Copperbelt schools when the question is about Mongu)
                sample_features = [{"_note": (
                    f"No data found within {_radius_km} km of {_location or 'the selected point'} "
                    f"in the pre-loaded offline dataset. The live GeoHub server is currently "
                    f"unavailable. Tell the user this clearly — do NOT describe data from "
                    f"other districts or provinces."
                )}]
                map_geojson = {"type": "FeatureCollection", "features": []}

    # If cross-context has settlement points and the current map has no Point features,
    # override with settlement points so the map always shows something useful.
    if _cross_context.get("settlement_geojson"):
        _has_points = map_geojson and any(
            (f.get("geometry") or {}).get("type") == "Point"
            for f in map_geojson.get("features", [])
        )
        if not _has_points:
            map_geojson = _cross_context["settlement_geojson"]

    # If cross-context has road lines, inject them as a context layer so they appear on the map.
    _ctx_layers = list(_CONTEXT_LAYERS)
    if _cross_context.get("road_geojson"):
        _ctx_layers = _ctx_layers + [
            {"geojson": _cross_context["road_geojson"], "name": "Roads", "type": "road"}
        ]

    ds = datasets[0] if datasets else {}

    # --- MAP IMAGE ANALYSIS ---
    if intent == "map_analysis":
        _img_b64 = st.session_state.get("uploaded_img_b64", "")
        _img_mime = st.session_state.get("uploaded_img_mime", "image/jpeg")
        _img_name = st.session_state.get("uploaded_img_name", "Uploaded map")

        with st.chat_message("assistant"):
            st.markdown('<span class="intent-badge intent-report">Map Analysis</span>', unsafe_allow_html=True)

            # Show the uploaded image inline so the user sees what's being analysed
            import base64 as _b64img
            _img_bytes_display = _b64img.b64decode(_img_b64)
            st.image(_img_bytes_display, caption=_img_name, use_container_width=True)

            _analysis_messages = [
                {"role": "user", "content": [
                    {"type": "image", "source": {
                        "type": "base64", "media_type": _img_mime, "data": _img_b64,
                    }},
                    {"type": "text", "text": map_analysis_user_prompt(question, _img_name)},
                ]}
            ]

            st.session_state.stop_streaming = False
            st.session_state.is_generating = True
            _sess_buf_key_ma = id(st.session_state) + 1
            _STREAM_BUFFERS[_sess_buf_key_ma] = ""

            def _map_analysis_stream():
                for chunk in claude.stream_with_history(map_analysis_system_prompt(), _analysis_messages, max_tokens=2000):
                    if st.session_state.get("stop_streaming"):
                        break
                    _STREAM_BUFFERS[_sess_buf_key_ma] = _STREAM_BUFFERS.get(_sess_buf_key_ma, "") + chunk
                    yield chunk

            try:
                analysis_text = st.write_stream(_map_analysis_stream())
            except Exception:
                analysis_text = _STREAM_BUFFERS.get(_sess_buf_key_ma, "⚠️ Something went wrong. Please try again.")
                st.warning(analysis_text)
            finally:
                st.session_state.is_generating = False
                st.session_state.stop_streaming = False
                _STREAM_BUFFERS.pop(_sess_buf_key_ma, None)

            # Export buttons — ask the user which format they want
            st.markdown("---")
            st.markdown("**Export this analysis:**")
            _exp_col1, _exp_col2, _exp_col3 = st.columns(3)

            with st.spinner("Preparing exports..."):
                _exp_docx = builder.to_docx(_img_name, analysis_text, {"name": _img_name, "description": f"Map analysis of {_img_name}"})
                _exp_pdf = builder.to_pdf(_img_name, analysis_text, {"name": _img_name, "description": f"Map analysis of {_img_name}"})

            _safe_name = _img_name.rsplit(".", 1)[0].replace(" ", "_")
            _exp_col1.download_button(
                "⬇️ Word (.docx)", _exp_docx,
                file_name=f"{_safe_name}_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"dl_map_docx_{len(st.session_state.messages)}",
                use_container_width=True,
            )
            _exp_col2.download_button(
                "⬇️ PDF", _exp_pdf,
                file_name=f"{_safe_name}_analysis.pdf",
                mime="application/pdf",
                key=f"dl_map_pdf_{len(st.session_state.messages)}",
                use_container_width=True,
            )
            # Table export — convert the markdown bullet points to a CSV-like text
            import pandas as _pd_map
            _table_rows = [
                {"Section": ln.lstrip("#").strip(), "Content": ""}
                if ln.startswith("#") else
                {"Section": "", "Content": ln.lstrip("•- ").strip()}
                for ln in analysis_text.splitlines()
                if ln.strip()
            ]
            _table_csv = "\n".join(
                f"{r['Section']}\t{r['Content']}" for r in _table_rows
            )
            _exp_col3.download_button(
                "⬇️ Table (.txt)", _table_csv,
                file_name=f"{_safe_name}_analysis_table.txt",
                mime="text/plain",
                key=f"dl_map_tbl_{len(st.session_state.messages)}",
                use_container_width=True,
            )

            st.session_state.messages.append({
                "role": "assistant", "content": analysis_text, "intent": intent,
                "ds_name": _img_name, "geojson": None,
                "docx_bytes": _exp_docx, "pdf_bytes": _exp_pdf,
            })
        return

    # --- REPORT ---
    if intent == "report":
        with st.chat_message("assistant"):
            st.markdown('<span class="intent-badge intent-report">Report</span>', unsafe_allow_html=True)
            if not ds:
                response = "I could not find a matching dataset on the Zambia GeoHub for your report request. Please try a more specific dataset name."
                st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response, "intent": intent})
                return

            with st.spinner("Generating report (~15 seconds)..."):
                stats = summarize_geojson(map_geojson) if map_geojson else {"feature_count": 0, "geometry_type": "Unknown", "fields": [], "numeric_stats": {}, "exceeded_limit": False}
                _lang_sfx = _LANG_INSTRUCTIONS.get(st.session_state.get("_lang", "English"), "")
                rpt_text = claude.ask(
                    system=report_system_prompt() + _lang_sfx,
                    user=report_prompt(ds["name"], ds["description"], ds.get("fields", []), stats, sample_features),
                    max_tokens=3000,
                )

            with st.spinner("Building Word and PDF..."):
                docx_bytes = builder.to_docx(ds["name"], rpt_text, ds)
                pdf_bytes = builder.to_pdf(ds["name"], rpt_text, ds)

            st.markdown(f"**Report: {ds['name']}**")
            st.markdown(rpt_text)
            st.markdown("**Download report:**")
            c1, c2 = st.columns(2)
            c1.download_button("⬇️ Word (.docx)", docx_bytes,
                file_name=f"{ds['name'].replace(' ','_')}_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_docx_new", use_container_width=True)
            c2.download_button("⬇️ PDF", pdf_bytes,
                file_name=f"{ds['name'].replace(' ','_')}_report.pdf",
                mime="application/pdf", key="dl_pdf_new", use_container_width=True)

            display_geojson = map_geojson or {"type": "FeatureCollection", "features": []}
            st_folium(_map(display_geojson, ds["name"], with_context=_is_point_geojson(display_geojson), highlight_location=_location or "", draw_bbox=_draw_bbox), width=720, height=340, returned_objects=[], key="map_new_rpt")

            st.session_state.messages.append({
                "role": "assistant", "content": rpt_text, "intent": intent,
                "docx_bytes": docx_bytes, "pdf_bytes": pdf_bytes,
                "ds_name": ds["name"], "geojson": map_geojson,
                "draw_bbox": _draw_bbox,
            })

    # --- SUMMARY ---
    elif intent == "summary":
        with st.chat_message("assistant"):
            st.markdown('<span class="intent-badge intent-summary">Summary</span>', unsafe_allow_html=True)
            if not ds:
                response = "I could not find a matching dataset to summarise. Try a more specific name like 'summarise health facilities'."
                st.markdown(response)
                st.session_state.messages.append({"role": "assistant", "content": response, "intent": intent})
                return

            stats = summarize_geojson(map_geojson) if map_geojson else {"feature_count": 0, "geometry_type": "Unknown", "fields": [], "numeric_stats": {}, "exceeded_limit": False}
            with st.spinner("Generating summary..."):
                summary = claude.ask(
                    system=summarizer_system_prompt(),
                    user=summarizer_prompt(ds["name"], ds["description"], ds.get("fields", []), sample_features, stats["feature_count"]),
                    max_tokens=1024,
                )

            if map_geojson:
                s = stats
                c1, c2, c3 = st.columns(3)
                c1.metric("Features", f"{s['feature_count']:,}")
                c2.metric("Geometry", s["geometry_type"].replace("esriGeometry", ""))
                c3.metric("Fields", len(s["fields"]))

            _render_data_tables(sample_features, ds["name"], key_prefix="new_sum")

            st.markdown(f"**Summary: {ds['name']}**")
            st.markdown(summary)
            st.download_button("⬇️ Download Summary (.txt)", summary,
                file_name=f"{ds['name'].replace(' ','_')}_summary.txt",
                mime="text/plain", key="dl_sum_new")

            display_geojson = map_geojson or {"type": "FeatureCollection", "features": []}
            st_folium(_map(display_geojson, ds["name"], with_context=_is_point_geojson(display_geojson), highlight_location=_location or "", draw_bbox=_draw_bbox), width=720, height=340, returned_objects=[], key="map_new_sum")

            st.session_state.messages.append({
                "role": "assistant", "content": summary, "intent": intent,
                "summary_txt": summary, "ds_name": ds["name"], "geojson": map_geojson,
                "sample_features": sample_features,
                "draw_bbox": _draw_bbox,
            })

    # --- CHAT (default) ---
    else:
        with st.chat_message("assistant"):
            st.markdown('<span class="intent-badge intent-chat">Answer</span>', unsafe_allow_html=True)

            # District / Province overview card — shown when the question names a location,
            # giving the user context about the area even if they're unfamiliar with it.
            if _location and _loc_type in ("district", "province"):
                _ov_key = f"ov_{_location}_{len(st.session_state.messages)}"
                with st.spinner(f"Loading {_location} overview..."):
                    _render_location_overview(_location, _loc_type, hub, key_prefix=_ov_key)

            # Build multi-turn message history for Claude so follow-up questions
            # reference previous answers (e.g. "how many of those are in Lusaka?")
            history = []
            for m in st.session_state.messages[:-1]:  # exclude the just-added user msg
                if m["role"] in ("user", "assistant") and m.get("content"):
                    history.append({"role": m["role"], "content": m["content"]})
            # Add current user prompt (with dataset context) as the final user turn
            _compare_note = (
                f"\n⚡ COMPARISON REQUEST: The user wants a side-by-side comparison of "
                f"{' and '.join(_comparison_locs)}. Use the sample records to compare "
                f"counts, types, and coverage between these locations. Present as a comparison table if possible."
                if _comparison_locs else ""
            )
            # Append uploaded document context if present
            _doc_ctx = ""
            _doc_text = st.session_state.get("uploaded_doc_text", "")
            _doc_name = st.session_state.get("uploaded_doc_name", "")
            if _doc_text:
                _doc_ctx = (
                    f"\n\n--- UPLOADED DOCUMENT: {_doc_name} ---\n"
                    f"{_doc_text[:6000]}"  # cap at 6000 chars to stay within token limits
                    f"\n--- END OF DOCUMENT ---\n"
                    "Use the above document as additional context when answering the question. "
                    "If the question is specifically about the document, prioritise its content."
                )
            _draw_counts = st.session_state.get("_draw_counts", {})
            _draw_counts_note = ""
            if _draw_counts:
                _draw_counts_note = " Feature counts in this area: " + ", ".join(
                    f"{l.split(' ',1)[-1]}: {c}" for l, c in _draw_counts.items()
                ) + "."
            _bbox_note = (
                f"\n[The user drew an area on the map. {_draw_bbox_location_note}. "
                f"Coordinates: lat {_draw_bbox['min_lat']:.3f}–{_draw_bbox['max_lat']:.3f}, "
                f"lon {_draw_bbox['min_lon']:.3f}–{_draw_bbox['max_lon']:.3f}. "
                f"Measurement: {_draw_bbox.get('measurement', 'N/A')}."
                f"{_draw_counts_note} "
                f"All data shown is filtered to this area. When answering 'where is this', "
                f"state the district and province identified above.]"
                if _draw_bbox else ""
            )
            user_p = chatbot_user_prompt(question + _compare_note + _doc_ctx + _bbox_note, datasets, sample_features, all_catalog=hub.get_catalog(), total_count=_total_count, location=_location or "", cross_context=_cross_context)

            # If a map image is attached, send it as a vision message block
            _img_b64 = st.session_state.get("uploaded_img_b64", "")
            _img_mime = st.session_state.get("uploaded_img_mime", "image/jpeg")
            _img_name = st.session_state.get("uploaded_img_name", "")
            if _img_b64:
                history.append({"role": "user", "content": [
                    {"type": "image", "source": {
                        "type": "base64", "media_type": _img_mime, "data": _img_b64,
                    }},
                    {"type": "text", "text": (
                        f"The user has uploaded a map image ({_img_name}). "
                        "Describe what geographic area, features, or patterns you can see in this image. "
                        "Then answer their question using both the image and the GeoHub dataset information provided.\n\n"
                        + user_p
                    )},
                ]})
            else:
                history.append({"role": "user", "content": user_p})

            # --- Streaming with stop button ---
            st.session_state.stop_streaming = False
            st.session_state.is_generating = True

            # Per-session buffer key (stable within a session)
            _sess_buf_key = id(st.session_state)
            _STREAM_BUFFERS[_sess_buf_key] = ""

            _lang_suffix = _LANG_INSTRUCTIONS.get(st.session_state.get("_lang", "English"), "")
            _chat_system = chatbot_system_prompt() + _lang_suffix

            def _stoppable_stream():
                for chunk in claude.stream_with_history(_chat_system, history, max_tokens=1500):
                    if st.session_state.get("stop_streaming"):
                        break
                    _STREAM_BUFFERS[_sess_buf_key] = _STREAM_BUFFERS.get(_sess_buf_key, "") + chunk
                    yield chunk

            _ai_error = False
            try:
                response = st.write_stream(_stoppable_stream())
            except Exception as e:
                # Recover any partial text already streamed
                response = _STREAM_BUFFERS.get(_sess_buf_key, "")
                _err_str = str(e)
                if not response:
                    if "overloaded" in _err_str.lower():
                        response = "⚠️ The AI is temporarily overloaded. Please try again in a few seconds."
                    elif "rate_limit" in _err_str.lower():
                        response = "⚠️ Rate limit reached. Please wait a moment and try again."
                    else:
                        response = "⚠️ Something went wrong with the AI response. Please try again."
                    _ai_error = True
                    st.warning(response)
            finally:
                st.session_state.is_generating = False
                st.session_state.stop_streaming = False
                _STREAM_BUFFERS.pop(_sess_buf_key, None)

            if not _ai_error:
                if datasets:
                    with st.expander("📂 Data sources"):
                        for d in datasets:
                            st.markdown(f"- **{d['name']}**  \n  {d['description'][:180]}")
                        st.markdown("[Browse all datasets on Zambia GeoHub ↗](https://zmb-geowb.hub.arcgis.com/search?collection=dataset&tags=zmb)")

            # Track live vs static fetch stats
            _fs = st.session_state.get("_fetch_stats", {"live": 0, "static": 0})
            if geojson and not _ai_error:
                if st.session_state.get("_last_fetch_was_live", False):
                    _fs["live"] += 1
                else:
                    _fs["static"] += 1
            st.session_state["_fetch_stats"] = _fs

            # Append message first, then show on-demand panel using the stored message
            _new_msg = {
                "role": "assistant", "content": response, "intent": intent,
                "ds_name": ds.get("name", ""), "geojson": map_geojson if not _ai_error else None,
                "location": _location or "",
                "sample_features": sample_features if not _ai_error else [],
                "buffer_center": _buffer_center,
                "buffer_radius_km": _radius_km,
                "draw_bbox": _draw_bbox,
                "data_source_url": ds.get("url", "") if not _ai_error else "",
                "data_live": st.session_state.get("_last_fetch_was_live", False),
            }
            st.session_state.messages.append(_new_msg)

            # Show map/table/chart buttons for this fresh answer
            if not _ai_error:
                _render_ondemand_panel(len(st.session_state.messages) - 1, _new_msg, ctx_layers=_ctx_layers)

# ---------------------------------------------------------------------------
# Handle pending edited question
# ---------------------------------------------------------------------------
if hasattr(st.session_state, "_pending_question") and st.session_state._pending_question:
    q = st.session_state._pending_question
    st.session_state._pending_question = None
    st.session_state.messages.append({"role": "user", "content": q})
    with st.chat_message("user"):
        st.markdown(q)
    process_question(q)
    st.rerun()



# ---------------------------------------------------------------------------
# Chat input — paperclip built into bar via accept_file
# ---------------------------------------------------------------------------
_chat_placeholder = (
    "Ask a question about your map — AI will analyse and extract information from it..."
    if st.session_state.get("uploaded_img_b64")
    else "Ask a question, say 'generate a report on...', or 'summarise...'"
)
_chat_result = st.chat_input(_chat_placeholder, accept_file="multiple",
                              file_type=["pdf", "docx", "txt", "png", "jpg", "jpeg", "webp"])
if _chat_result:
    for _af in (_chat_result.files or []):
        _af_name = _af.name.lower()
        if _af_name.endswith((".png", ".jpg", ".jpeg", ".webp")):
            import base64 as _b64
            _img_bytes = _af.read()
            _img_mime = "image/png" if _af_name.endswith(".png") else "image/jpeg"
            st.session_state["uploaded_img_b64"]  = _b64.b64encode(_img_bytes).decode()
            st.session_state["uploaded_img_mime"] = _img_mime
            st.session_state["uploaded_img_name"] = _af.name
        else:
            try:
                if _af_name.endswith(".pdf"):
                    import pypdf as _pypdf
                    _reader = _pypdf.PdfReader(_af)
                    _doc_text = "\n".join(p.extract_text() or "" for p in _reader.pages)
                elif _af_name.endswith(".docx"):
                    import docx as _docx
                    _doc_obj = _docx.Document(_af)
                    _doc_text = "\n".join(p.text for p in _doc_obj.paragraphs if p.text.strip())
                else:
                    _doc_text = _af.read().decode("utf-8", errors="ignore")
                if _doc_text.strip():
                    st.session_state["uploaded_doc_text"] = _doc_text.strip()
                    st.session_state["uploaded_doc_name"] = _af.name
            except Exception:
                pass
    question = _chat_result.text or ""
    if question.strip() or st.session_state.get("uploaded_img_b64") or st.session_state.get("uploaded_doc_text"):
        if st.session_state.edit_idx is None:
            if "_current_chat_id" not in st.session_state:
                import uuid as _uuid2
                st.session_state["_current_chat_id"] = str(_uuid2.uuid4())
            _display_q = question.strip() or f"[Attached: {(_chat_result.files or ['file'])[0].name if _chat_result.files else 'file'}]"
            st.session_state.messages.append({"role": "user", "content": _display_q})
            with st.chat_message("user"):
                st.markdown(_display_q)
            process_question(question)
            _save_current_chat()
            st.session_state["_scroll_to_bottom"] = True
            st.rerun()

# Auto-scroll to bottom after a new answer is generated
if st.session_state.pop("_scroll_to_bottom", False):
    st.components.v1.html("""
        <script>
            (function tryScroll(attempts) {
                var doc = window.parent.document;
                var el = doc.querySelector('[data-testid="stAppViewBlockContainer"]')
                    || doc.querySelector('section.main')
                    || doc.querySelector('.main');
                if (el) {
                    el.scrollTop = el.scrollHeight;
                } else if (attempts > 0) {
                    setTimeout(function() { tryScroll(attempts - 1); }, 200);
                }
            })(10);
        </script>
    """, height=1)

# Persist current chat to URL after every render so refresh restores it
_persist_chat()
