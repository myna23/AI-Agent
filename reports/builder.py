"""
Report generation — Word (.docx) and PDF.

Both methods accept the AI-generated markdown text and
return bytes that can be served directly via st.download_button().

Markdown conventions expected from the AI:
  ## Heading      → Heading 2
  ### Sub-heading → Heading 3
  - bullet        → List Bullet
  plain text      → Normal paragraph
"""

from __future__ import annotations

import re
from datetime import date
from io import BytesIO

# Word
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    HRFlowable,
)


class ReportBuilder:
    """Build Word and PDF reports from AI narrative + dataset metadata."""

    # ------------------------------------------------------------------
    # Word (.docx)
    # ------------------------------------------------------------------

    def to_docx(
        self,
        title: str,
        ai_content: str,
        dataset_meta: dict,
    ) -> bytes:
        """
        Build a Word document and return it as bytes.

        Parameters
        ----------
        title        : Report / dataset title
        ai_content   : AI markdown response text
        dataset_meta : Cleaned dataset dict from HubClient
        """
        doc = Document()

        # ---- Page margins ----
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # ---- Cover header ----
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0x1D, 0x35, 0x57)

        # Subtitle / metadata block
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"Zambia GeoHub AI Platform  •  Generated {date.today().strftime('%d %B %Y')}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(0x45, 0x7B, 0x9D)

        doc.add_paragraph()  # spacer

        # Dataset metadata table
        fields_preview = ", ".join(
            f["alias"] or f["name"] for f in dataset_meta.get("fields", [])[:8]
        )
        table_data = [
            ("Dataset", dataset_meta.get("name", title)),
            ("Geometry type", dataset_meta.get("geometry_type", "—")),
            ("Key fields", fields_preview or "—"),
            ("Source", "Zambia GeoHub — zmb-geowb.hub.arcgis.com"),
            ("Report date", date.today().strftime("%d %B %Y")),
        ]
        tbl = doc.add_table(rows=len(table_data), cols=2)
        tbl.style = "Light List Accent 1"
        for i, (label, value) in enumerate(table_data):
            tbl.rows[i].cells[0].text = label
            tbl.rows[i].cells[1].text = str(value)
            tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()  # spacer

        # ---- AI narrative ----
        self._parse_markdown_to_docx(doc, ai_content)

        # ---- Footer note ----
        doc.add_paragraph()
        footer = doc.add_paragraph(
            "This report was automatically generated using the Zambia GeoHub AI Platform "
            "powered by AI. Data sourced from zmb-geowb.hub.arcgis.com. "
            "Always verify against the latest official datasets before making decisions."
        )
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _parse_markdown_to_docx(self, doc: Document, md_text: str) -> None:
        """Walk the AI markdown output and map it to Word paragraph styles."""
        for line in md_text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            if stripped.startswith("### "):
                p = doc.add_heading(stripped[4:], level=3)
                p.runs[0].font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
            elif stripped.startswith("## "):
                p = doc.add_heading(stripped[3:], level=2)
                p.runs[0].font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
            elif stripped.startswith("# "):
                p = doc.add_heading(stripped[2:], level=1)
                p.runs[0].font.color.rgb = RGBColor(0x1D, 0x35, 0x57)
            elif stripped.startswith(("- ", "* ", "• ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            else:
                # Bold inline (**text**) → simple replacement
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
                doc.add_paragraph(clean)

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def to_pdf(
        self,
        title: str,
        ai_content: str,
        dataset_meta: dict,
    ) -> bytes:
        """
        Build a PDF document and return it as bytes.
        """
        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            rightMargin=2.5 * cm,
            leftMargin=2.5 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        navy = colors.HexColor("#1D3557")
        steel = colors.HexColor("#457B9D")

        # Custom styles
        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            textColor=navy,
            fontSize=22,
            spaceAfter=6,
        )
        subtitle_style = ParagraphStyle(
            "Subtitle",
            parent=styles["Normal"],
            textColor=steel,
            fontSize=10,
            spaceAfter=12,
            alignment=1,  # centre
        )
        h2_style = ParagraphStyle(
            "H2",
            parent=styles["Heading2"],
            textColor=navy,
            fontSize=13,
            spaceBefore=12,
            spaceAfter=4,
        )
        h3_style = ParagraphStyle(
            "H3",
            parent=styles["Heading3"],
            textColor=steel,
            fontSize=11,
            spaceBefore=8,
            spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "Bullet",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            leftIndent=20,
            bulletIndent=10,
            spaceAfter=3,
        )
        footer_style = ParagraphStyle(
            "Footer",
            parent=styles["Normal"],
            fontSize=8,
            textColor=colors.grey,
            spaceBefore=20,
        )

        story = []

        # Title + subtitle
        story.append(Paragraph(title, title_style))
        story.append(
            Paragraph(
                f"Zambia GeoHub AI Platform  •  Generated {date.today().strftime('%d %B %Y')}",
                subtitle_style,
            )
        )
        story.append(HRFlowable(width="100%", thickness=1, color=steel, spaceAfter=10))

        # Metadata table
        fields_preview = ", ".join(
            f["alias"] or f["name"] for f in dataset_meta.get("fields", [])[:8]
        ) or "—"
        meta_rows = [
            ["Dataset", dataset_meta.get("name", title)],
            ["Geometry type", dataset_meta.get("geometry_type", "—")],
            ["Key fields", fields_preview],
            ["Source", "zmb-geowb.hub.arcgis.com"],
            ["Report date", date.today().strftime("%d %B %Y")],
        ]
        tbl = Table(meta_rows, colWidths=[4 * cm, 12 * cm])
        tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EBF2FA")),
                    ("TEXTCOLOR", (0, 0), (0, -1), navy),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
                    ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#F8FBFD")]),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(tbl)
        story.append(Spacer(1, 0.5 * cm))

        # AI narrative
        for line in ai_content.splitlines():
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 0.2 * cm))
                continue

            if stripped.startswith("### "):
                story.append(Paragraph(stripped[4:], h3_style))
            elif stripped.startswith("## "):
                story.append(Paragraph(stripped[3:], h2_style))
            elif stripped.startswith("# "):
                story.append(Paragraph(stripped[2:], h2_style))
            elif stripped.startswith(("- ", "* ", "• ")):
                clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped[2:])
                story.append(Paragraph(f"• {clean}", bullet_style))
            elif re.match(r"^\d+\.\s", stripped):
                clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", re.sub(r"^\d+\.\s", "", stripped))
                story.append(Paragraph(clean, bullet_style))
            else:
                clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped)
                story.append(Paragraph(clean, body_style))

        # Footer
        story.append(
            Paragraph(
                "This report was automatically generated using the Zambia GeoHub AI Platform "
                "powered by AI. Always verify data against official sources.",
                footer_style,
            )
        )

        doc.build(story)
        return buf.getvalue()
